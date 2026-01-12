#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
等保完结单批量导出脚本
将 Word 文档中的等保完结单内容批量导出到 Excel 表格
支持一个项目多个系统的情况
"""

import re
import os
from docx import Document
import pandas as pd
from datetime import datetime
from pathlib import Path
from openpyxl.styles import Alignment


def extract_personnel_contribution(text):
    """
    从人员工作安排及贡献率文本中提取简化格式
    支持多个人员，每个人一行
    """
    if not text:
        return ''

    result = []
    text_combined = ' '.join(text.replace('\n', ' ').split())

    # 匹配所有"姓名：...（百分比）"的模式
    pattern = r'([^：:：\s]+)\s*[：:：][^（(]*[（(](\d+%)[)）]'
    matches = re.findall(pattern, text_combined)

    if matches:
        for name, percentage in matches:
            result.append(f"{name.strip()}{percentage.strip()}")
        return '\n'.join(result)

    return text[:20] if len(text) > 20 else text


def extract_project_basic_info(doc):
    """提取项目基本信息"""
    basic_info = {
        '部门': '软测部',
        '项目编号': '',
        '项目名称': '',
        '客户单位名称': '',
        '项目地点': '',
        '项目类型': '',
        '启动时间': ''
    }

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            row_text = ''.join(cells)

            # 项目编号 - 改进识别逻辑
            if not basic_info['项目编号']:
                for cell in cells:
                    cell_text = cell.strip()
                    # 匹配 QZX 开头的编号格式（如 QZXGC-202512004 或 QZX2025-0444）
                    if cell_text.startswith('QZX') and '项目编号' not in cell_text:
                        # 排除包含中文的情况
                        if not any('\u4e00' <= c <= '\u9fff' for c in cell_text):
                            basic_info['项目编号'] = cell_text
                            break

            # 项目名称
            if '项目名称' in row_text and not basic_info['项目名称']:
                for j, cell in enumerate(cells):
                    if '项目名称' in cell:
                        for k in range(j + 1, len(cells)):
                            candidate = cells[k].strip()
                            # 修改判断条件：长度>8且不包含项目编号
                            if candidate and candidate != '项目名称' and 'QZX' not in candidate:
                                basic_info['项目名称'] = candidate
                                break
                        break

            # 客户单位名称
            if '客户单位名称' in row_text and not basic_info['客户单位名称']:
                for j, cell in enumerate(cells):
                    if '客户单位名称' in cell:
                        for k in range(j + 1, len(cells)):
                            if cells[k] and ('公司' in cells[k] or '中心' in cells[k] or len(cells[k]) > 8):
                                basic_info['客户单位名称'] = cells[k].strip()
                                break
                        break

            # 项目地点
            if '项目地点' in row_text and not basic_info['项目地点']:
                for j, cell in enumerate(cells):
                    if '项目地点' in cell:
                        for k in range(j + 1, len(cells)):
                            candidate = cells[k].strip()
                            if candidate and candidate != '项目地点' and '、' not in candidate and len(candidate) <= 8:
                                basic_info['项目地点'] = candidate
                                break
                        break

            # 启动时间
            if '启动时间' in row_text and not basic_info['启动时间']:
                for j, cell in enumerate(cells):
                    if '启动时间' in cell:
                        for k in range(j + 1, len(cells)):
                            time_text = cells[k].strip()
                            if time_text and ('/' in time_text or '-' in time_text):
                                basic_info['启动时间'] = time_text
                                break
                        break

            # 项目类型 - 识别带●符号的选项
            if '项目类型' in row_text and not basic_info['项目类型']:
                # print(f"  [调试] 找到项目类型行，单元格内容: {cells}")

                # 定义项目类型关键词
                type_keywords = ['等保', '密评', '风评', '安评', '数评', '软测', '安服']

                # 在所有单元格中查找带●的项目类型
                for cell in cells:
                    if '●' in cell:
                        # print(f"  [调试] 找到●符号，单元格内容: '{cell}'")

                        # 检查●后面是否紧跟项目类型关键词（可能有空格）
                        for keyword in type_keywords:
                            # 匹配 "●关键词" 或 "● 关键词" 或 "●  关键词" 等模式
                            pattern = f'●\\s*{keyword}'
                            if re.search(pattern, cell):
                                basic_info['项目类型'] = keyword
                                # print(f"  [调试] 识别为：{keyword}")
                                break

                        if basic_info['项目类型']:
                            break

                # 如果还是没找到，尝试在整行文本中查找
                if not basic_info['项目类型']:
                    for keyword in type_keywords:
                        pattern = f'●\\s*{keyword}'
                        if re.search(pattern, row_text):
                            basic_info['项目类型'] = keyword
                            # print(f"  [调试] 从整行识别为：{keyword}")
                            break

    return basic_info


def extract_systems_info(doc):
    """提取系统信息"""
    systems = []

    for table in doc.tables:
        header_row_idx = -1
        header_mapping = {}

        for i, row in enumerate(table.rows):
            cells_text = [cell.text.strip() for cell in row.cells]
            row_text = ''.join(cells_text)

            if '系统名称' in row_text and '系统级别' in row_text and '系统类型' in row_text:
                header_row_idx = i

                for j, cell_text in enumerate(cells_text):
                    if '系统名称' in cell_text:
                        header_mapping['系统名称'] = j
                    elif '系统级别' in cell_text:
                        header_mapping['系统级别'] = j
                    elif '系统类型' in cell_text:
                        header_mapping['系统类型'] = j
                    elif '人员工作安排' in cell_text or '贡献率' in cell_text:
                        header_mapping['人员贡献率'] = j
                break

        if header_row_idx >= 0 and header_mapping:
            for i in range(header_row_idx + 1, len(table.rows)):
                row = table.rows[i]
                cells = [cell.text.strip() for cell in row.cells]

                if not cells or not cells[0] or '业务确认' in ''.join(cells) or '部门经理' in ''.join(cells):
                    continue

                system_info = {
                    '系统名称': '',
                    '系统级别': '',
                    '系统类型': '',
                    '人员贡献率': ''
                }

                if '系统名称' in header_mapping and header_mapping['系统名称'] < len(cells):
                    system_info['系统名称'] = cells[header_mapping['系统名称']].strip()

                if '系统级别' in header_mapping and header_mapping['系统级别'] < len(cells):
                    level = cells[header_mapping['系统级别']].strip()
                    if level.isdigit() or level == '/':
                        system_info['系统级别'] = level

                if '系统类型' in header_mapping and header_mapping['系统类型'] < len(cells):
                    system_info['系统类型'] = cells[header_mapping['系统类型']].strip()

                if '人员贡献率' in header_mapping and header_mapping['人员贡献率'] < len(cells):
                    raw_text = cells[header_mapping['人员贡献率']].strip()
                    system_info['人员贡献率'] = extract_personnel_contribution(raw_text)

                if system_info['系统名称']:
                    systems.append(system_info)

    return systems


def process_single_doc(docx_path, start_sequence):
    """处理单个文档"""
    try:
        doc = Document(docx_path)
        basic_info = extract_project_basic_info(doc)
        systems = extract_systems_info(doc)

        project_type = basic_info['项目类型']
        if project_type in ['风评', '安评', '数评', '软测', '安服']:
            for system in systems:
                system['系统级别'] = '/'
                system['系统类型'] = '/'

        project_records = []
        if systems:
            for idx, system in enumerate(systems):
                record = {
                    '序号': start_sequence + idx,
                    **basic_info,
                    **system
                }
                project_records.append(record)
        else:
            record = {
                '序号': start_sequence,
                **basic_info,
                '系统名称': '',
                '系统级别': '/' if project_type in ['风评', '安评', '数评', '软测', '安服'] else '',
                '系统类型': '/' if project_type in ['风评', '安评', '数评', '软测', '安服'] else '',
                '人员贡献率': ''
            }
            project_records.append(record)

        print(f"\n✓ 成功提取: {os.path.basename(docx_path)}")
        print(f"  项目名称: {basic_info['项目名称']}")
        print(f"  项目类型: {basic_info['项目类型']}")
        print(f"  系统数量: {len(systems)}")

        return project_records

    except Exception as e:
        print(f"\n✗ 处理失败: {os.path.basename(docx_path)}")
        print(f"  错误: {str(e)}")
        return []


def get_quarter_from_date(date_str):
    """根据日期确定季度"""
    if not date_str:
        return 1, datetime.now().year

    try:
        if '/' in date_str:
            date = datetime.strptime(date_str, '%Y/%m/%d')
        elif '-' in date_str:
            date = datetime.strptime(date_str, '%Y-%m-%d')
        else:
            return 1, datetime.now().year

        month = date.month
        year = date.year

        if month <= 3:
            quarter = 1
        elif month <= 6:
            quarter = 2
        elif month <= 9:
            quarter = 3
        else:
            quarter = 4

        return quarter, year
    except:
        return 1, datetime.now().year


def batch_process_docs(input_folder='.', pattern='*.docx'):
    """批量处理文档"""
    folder_path = Path(input_folder)
    docx_files = list(folder_path.glob(pattern))

    # 过滤文件
    filtered_files = []
    for f in docx_files:
        if f.name.startswith('~$'):
            continue
        if '年第' in f.name and '季度项目完结单' in f.name:
            continue
        filtered_files.append(f)

    docx_files = filtered_files

    if not docx_files:
        print(f"未找到文件")
        return [], None, None

    print(f"找到 {len(docx_files)} 个文档:")
    for f in docx_files:
        print(f"  - {f.name}")

    print("\n开始处理...")

    all_records = []
    current_sequence = 1
    earliest_quarter = None
    earliest_year = None

    for docx_file in docx_files:
        records = process_single_doc(str(docx_file), current_sequence)
        if records:
            all_records.extend(records)
            current_sequence += len(records)

            if records[0].get('启动时间'):
                quarter, year = get_quarter_from_date(records[0]['启动时间'])
                if earliest_quarter is None or (year < earliest_year) or (
                        year == earliest_year and quarter < earliest_quarter):
                    earliest_quarter = quarter
                    earliest_year = year

    print(f"\n共提取 {len(all_records)} 条记录")

    if earliest_quarter is None:
        earliest_quarter, earliest_year = get_quarter_from_date(None)

    return all_records, earliest_quarter, earliest_year


def export_to_excel(project_list, output_path, quarter, year):
    """导出到Excel"""
    if not project_list:
        print("没有数据")
        return

    sheet_name = f'{year}年第{quarter}季度项目完结单'

    columns = [
        '序号', '部门', '项目编号', '项目名称', '客户单位名称',
        '项目地点', '项目类型', '系统名称', '系统级别', '系统类型', '人员贡献率'
    ]

    df = pd.DataFrame(project_list, columns=columns)

    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name=sheet_name, index=False)

        worksheet = writer.sheets[sheet_name]

        column_widths = {
            'A': 8, 'B': 10, 'C': 18, 'D': 45, 'E': 25,
            'F': 12, 'G': 10, 'H': 50, 'I': 10, 'J': 20, 'K': 65
        }

        for col, width in column_widths.items():
            worksheet.column_dimensions[col].width = width

        for row in worksheet.iter_rows(min_row=1, max_row=worksheet.max_row,
                                       min_col=1, max_col=len(columns)):
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        worksheet.row_dimensions[1].height = 25

    print(f"\n✓ 导出成功: {output_path}")
    print(f"  工作表: {sheet_name}")
    print(f"  记录数: {len(project_list)}")


def main():
    """主函数"""
    print("=" * 70)
    print("完结单批量导出工具")
    print("=" * 70)

    input_folder = "."
    file_pattern = "*完结单*.docx"

    print(f"\n配置:")
    print(f"  文件夹: {input_folder}")
    print(f"  模式: {file_pattern}\n")

    try:
        project_list, quarter, year = batch_process_docs(input_folder, file_pattern)

        if project_list:
            output_file = f"{year}年第{quarter}季度项目完结单.xlsx"

            print(f"\n季度: {year}年第{quarter}季度")
            print(f"输出: {output_file}")

            print("\n" + "=" * 70)
            print("数据摘要:")
            print("=" * 70)
            for proj in project_list:
                sys_name = proj.get('系统名称', '')
                sys_short = sys_name[:30] + '...' if len(sys_name) > 30 else sys_name
                print(f"{proj['序号']}. {proj['项目名称'][:40]} - {proj['项目类型']} - {sys_short}")

            print("\n" + "=" * 70)
            export_to_excel(project_list, output_file, quarter, year)
            print("=" * 70)
        else:
            print("\n未提取到数据")

    except Exception as e:
        print(f"\n错误: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()

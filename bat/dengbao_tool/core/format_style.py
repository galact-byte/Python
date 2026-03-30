"""文档格式定义 — 基于旧文档格式基线"""

from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH


# 页面设置
PAGE_WIDTH = Cm(21.0)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.18)

# 字体定义
FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_FANG = "仿宋"
FONT_FANG_GB = "仿宋_GB2312"

# 字号
SIZE_TITLE = Pt(22)        # 文档标题
SIZE_HEADING = Pt(14)      # 章节标题 / 小节标题 / 正文
SIZE_TABLE = Pt(12)        # 表格单元格

# 首行缩进
INDENT_2CHAR = Emu(355600)  # 约2个字符

# 行距
LINE_SPACING = Pt(28)       # 固定值28磅（公文标准）


def apply_title_format(paragraph):
    """应用文档标题格式：宋体22pt居中"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = FONT_SONG
        run.font.size = SIZE_TITLE
        run.font.bold = None
        _set_east_asia_font(run, FONT_SONG)


def apply_heading_format(paragraph):
    """应用章节标题格式：黑体14pt左对齐"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    for run in paragraph.runs:
        run.font.name = FONT_HEI
        run.font.size = SIZE_HEADING
        run.font.bold = None
        _set_east_asia_font(run, FONT_HEI)


def apply_subheading_format(paragraph):
    """应用小节标题格式：仿宋_GB2312 14pt 首行缩进"""
    paragraph.paragraph_format.first_line_indent = INDENT_2CHAR
    for run in paragraph.runs:
        run.font.name = FONT_FANG_GB
        run.font.size = SIZE_HEADING
        run.font.bold = None
        _set_east_asia_font(run, FONT_FANG_GB)


def apply_body_format(paragraph):
    """应用正文格式：仿宋 14pt 首行缩进"""
    paragraph.paragraph_format.first_line_indent = INDENT_2CHAR
    for run in paragraph.runs:
        run.font.name = FONT_FANG
        run.font.size = SIZE_HEADING
        run.font.bold = None
        _set_east_asia_font(run, FONT_FANG)


def apply_table_cell_format(cell):
    """应用表格单元格格式：仿宋_GB2312 12pt"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_FANG_GB
            run.font.size = SIZE_TABLE
            _set_east_asia_font(run, FONT_FANG_GB)


def set_cell_text(cell, text, font_name=FONT_FANG_GB, font_size=SIZE_TABLE):
    """设置单元格文本并应用格式"""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            _set_east_asia_font(run, font_name)


def add_formatted_paragraph(doc, text, fmt_type="body"):
    """添加带格式的段落"""
    p = doc.add_paragraph(text)
    if fmt_type == "title":
        apply_title_format(p)
    elif fmt_type == "heading":
        apply_heading_format(p)
    elif fmt_type == "subheading":
        apply_subheading_format(p)
    else:
        apply_body_format(p)
    return p


def _set_east_asia_font(run, font_name):
    """设置东亚字体（中文字体需要单独设置）"""
    from docx.oxml.ns import qn
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from lxml import etree
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
    rfonts.set(qn('w:eastAsia'), font_name)

"""新文档生成 — 精准 run 级填充，保留模版格式"""

import os
import shutil
import re
from copy import deepcopy
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from models.project_data import ProjectData, ReportInfo
from core.format_style import (
    set_cell_text, FONT_FANG_GB, FONT_SONG, SIZE_TABLE, _set_east_asia_font
)


# 生成态：备案表强制五号(10.5pt)；定级报告保留模板字号
_FILL_FORCE_SIZE = None


def _set_fill_mode(mode):
    """mode: 'beian' → 强制 10.5pt；'report' → 保留模板"""
    global _FILL_FORCE_SIZE
    if mode == 'beian':
        _FILL_FORCE_SIZE = Pt(10.5)
    else:
        _FILL_FORCE_SIZE = None


# ══════════════════════════════════════════════════════
#  底层工具函数
# ══════════════════════════════════════════════════════

def _highlight_run(run, color="yellow"):
    """给 run 添加高亮色"""
    rpr = run._element.get_or_add_rPr()
    for node in list(rpr):
        if node.tag == qn('w:highlight'):
            rpr.remove(node)
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), color)
    rpr.append(highlight)


def _highlight_cell(cell, color="yellow"):
    """给单元格中所有 run 添加高亮"""
    for para in cell.paragraphs:
        for run in para.runs:
            _highlight_run(run, color)


def _apply_fill_style(run):
    """统一填写内容样式为宋体五号。"""
    if run is None:
        return
    run.font.name = FONT_SONG
    run.font.size = Pt(10.5)
    _set_east_asia_font(run, FONT_SONG)


def _set_run_font_slots(run, east_asia_font=None, western_font=None):
    if run is None:
        return
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    if western_font:
        run.font.name = western_font
        rfonts.set(qn('w:ascii'), western_font)
        rfonts.set(qn('w:hAnsi'), western_font)
        rfonts.set(qn('w:cs'), western_font)
    if east_asia_font:
        rfonts.set(qn('w:eastAsia'), east_asia_font)


def _set_run_text(run, text):
    """\u586b\u5199\u503c\u7684\u5b57\u4f53\u7edf\u4e00\uff1a\u4e2d\u6587 \u4eff\u5b8b_GB2312\u3001\u897f\u6587 Times New Roman\uff0c\u5907\u6848\u8868 mode \u5f3a\u5236 10.5pt\u3002"""
    run.text = text
    raw_text = str(text or "")
    if raw_text:
        _set_run_font_slots(
            run,
            east_asia_font=FONT_FANG_GB,
            western_font="Times New Roman",
        )
    if _FILL_FORCE_SIZE is not None:
        run.font.size = _FILL_FORCE_SIZE
    elif run.font.size is None:
        run.font.size = Pt(10.5)


def _fill_text_into_underscores(text, value, min_tail=0):
    value = str(value or "")
    if "_" not in text:
        return value
    return re.sub(r"_+", value, text, count=1)


def _is_placeholder_text(text):
    if text is None:
        return False
    stripped = text.strip()
    return (not stripped) or all(ch in " _\u3000" for ch in text) or "_" in text


def _parse_date_parts(value):
    raw = str(value or "").strip()
    if not raw:
        return "", "", ""
    if re.match(r"^\d{4}-\d{2}-\d{2}$", raw):
        year, month, day = raw.split("-")
        return year, str(int(month)), str(int(day))
    match = re.search(r"(\d{4})\D+(\d{1,2})\D+(\d{1,2})", raw)
    if match:
        year, month, day = match.groups()
        return year, str(int(month)), str(int(day))
    return raw, "", ""


def _fill_date_cell(cell, value):
    year, month, day = _parse_date_parts(value)
    if not year:
        return False
    parts = [year, month, day]
    part_idx = 0
    for para in cell.paragraphs:
        runs = para.runs
        for idx, run in enumerate(runs):
            if any(mark in run.text for mark in ("年", "月", "日")) and idx > 0:
                prev_run = runs[idx - 1]
                if _is_placeholder_text(prev_run.text) and part_idx < len(parts):
                    _set_run_text(prev_run, parts[part_idx])
                    part_idx += 1
        if part_idx >= len(parts):
            return True
    return part_idx > 0


def _check_first_sym_in_paragraph(paragraph, checked):
    if paragraph is None:
        return
    for run in paragraph.runs:
        if run._element.find(qn('w:sym')) is not None:
            _check_sym(run, checked)
            return


# ══════════════════════════════════════════════════════
#  底层工具函数
# ══════════════════════════════════════════════════════

def _check_sym(run, checked=True):
    """将 run 中的 w:sym 元素设为勾选或未勾选"""
    sym = run._element.find(qn('w:sym'))
    if sym is None:
        return
    if checked:
        sym.set(qn('w:font'), 'Wingdings')
        sym.set(qn('w:char'), 'F0FE')
    else:
        # 恢复为未勾选（保持原始字体）
        current_font = sym.get(qn('w:font'))
        if current_font == 'Wingdings':
            sym.set(qn('w:font'), 'Wingdings 2')
        sym.set(qn('w:char'), '0030')


def _find_and_check(cell, match_text, multi=False):
    """
    在单元格中查找 match_text 对应选项并打勾。
    match_text 可以是选项编号（如"4"）或选项文字（如"企业"）。
    multi=True 时可同时勾选多个匹配项，返回是否至少命中一次。

    结构: [sym_run][编号_run][文字_run][空格_run] [sym_run][编号_run]...
    sym_run 在编号/文字 run 的前面。
    """
    matched = False
    for para in cell.paragraphs:
        runs = para.runs
        for i, run in enumerate(runs):
            sym = run._element.find(qn('w:sym'))
            if sym is None:
                continue
            # 找 sym 后面的选项编号和文字
            option_parts = []
            for j in range(i + 1, min(i + 4, len(runs))):
                next_run = runs[j]
                # 碰到下一个 sym 就停
                if next_run._element.find(qn('w:sym')) is not None:
                    break
                text = next_run.text.strip()
                if text:
                    option_parts.append(text)

            option_str = ''.join(option_parts)
            # 匹配：编号匹配、文字匹配、或编号+文字匹配
            if (match_text in option_parts or
                match_text in option_str or
                any(match_text == p for p in option_parts)):
                _check_sym(run, True)
                matched = True
                if not multi:
                    return True
    return matched


def _fill_scenario_options(cell, text):
    """场景行（多选 + 其他___）：尝试匹配每个 token 勾选；未命中的 token 写入 其他___。"""
    if not text:
        return
    raw = str(text).replace('，', ',').replace('、', ',').replace(';', ',').replace('；', ',')
    tokens = [t.strip() for t in raw.split(',') if t.strip()]
    leftover = []
    for token in tokens:
        if not _find_and_check(cell, token, multi=True):
            leftover.append(token)
    if leftover:
        _find_and_check(cell, '其他', multi=True)
        _fill_other_option(cell, '其他', '、'.join(leftover))


def _find_and_check_multiple(cell, match_texts):
    """勾选多个选项"""
    for text in match_texts:
        if text:
            _find_and_check(cell, text, multi=True)


def _fill_run(cell, para_idx, run_idx, text):
    """精准替换指定位置 run 的文本，不动其他 run"""
    try:
        para = cell.paragraphs[para_idx]
        run = para.runs[run_idx]
        run.text = text
    except (IndexError, AttributeError):
        pass


def _fill_after_keyword(cell, keyword, text):
    """在单元格中找到 keyword 后面的空白/下划线 run，填入 text"""
    for para in cell.paragraphs:
        runs = para.runs
        found_keyword = False
        for i, run in enumerate(runs):
            if keyword in run.text:
                found_keyword = True
                if "_" in run.text and re.search(r"_+", run.text):
                    if _fill_placeholder_run(run, text):
                        return True
                continue
            if found_keyword and (run.text.strip() == '' or
                                   all(c in ' _\u3000' for c in run.text)):
                if '_' in run.text:
                    _fill_placeholder_run(run, text)
                else:
                    _set_run_text(run, '  ' + text)
                return True
    return False


def _copy_run_style(src_run, dst_run):
    """复制 run 的基础字体样式。"""
    if src_run is None or dst_run is None:
        return
    src_rpr = src_run._element.rPr
    if src_rpr is not None:
        dst_rpr = dst_run._element.rPr
        if dst_rpr is not None:
            dst_run._element.remove(dst_rpr)
        dst_run._element.insert(0, deepcopy(src_rpr))
    # 注意：不要再用 dst_run.font.name = src_run.font.name 等高层 API 二次覆盖。
    # 当 src.font.name=None（如模板只有 eastAsia 没 ascii）时，setter 会删除整个 rFonts，
    # 一并清掉 eastAsia 仿宋_GB2312，导致中文回退到默认宋体。
    # rPr 的 deepcopy 已经包含 font/size/bold/italic 完整信息。
    if src_run.font.name and dst_run.font.name != src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.size is not None and dst_run.font.size != src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.bold is not None and dst_run.font.bold != src_run.font.bold:
        dst_run.font.bold = src_run.font.bold
    if src_run.font.italic is not None and dst_run.font.italic != src_run.font.italic:
        dst_run.font.italic = src_run.font.italic


def _insert_run_after(run):
    """在当前 run 后插入一个新 run，并保留段落归属。"""
    if run is None:
        return None
    paragraph = run._parent
    new_run = paragraph.add_run("")
    run._element.addnext(new_run._element)
    return new_run


def _fill_placeholder_run(run, value, pattern=r"_+", keep_length=False):
    """
    仅替换当前 run 中命中的下划线占位段，并让填充值本身保留下划线效果。
    不再把 `_` 字符残留在填充值后面。

    keep_length=True：保留原占位"视觉长度"——value 短于原下划线时，在 value 之后追加
    剩余 `_` 字符，保持视觉填空线效果（适用于表六数据来源/流出/存储位置等需要视觉占位的字段）。
    """
    if run is None:
        return False
    value = str(value or "")
    if not value:
        return False
    original_text = run.text or ""
    match = re.search(pattern, original_text)
    if not match:
        return False

    prefix = original_text[:match.start()]
    suffix = original_text[match.end():]

    # 计算剩余下划线（保持原视觉长度）
    extra_underscores = ""
    if keep_length:
        underline_count = match.end() - match.start()
        value_visual = sum(2 if '一' <= c <= '鿿' else 1 for c in value)
        remaining = underline_count - value_visual
        if remaining > 0:
            extra_underscores = "_" * remaining

    if prefix:
        fill_run = _insert_run_after(run)
        if fill_run is None:
            return False
        _copy_run_style(run, fill_run)
        suffix_run = None
        if suffix or extra_underscores:
            suffix_run = _insert_run_after(fill_run)
            if suffix_run is None:
                return False
            _copy_run_style(run, suffix_run)
        _set_run_text(run, prefix)
    else:
        fill_run = run
        suffix_run = None
        if suffix or extra_underscores:
            suffix_run = _insert_run_after(run)
            if suffix_run is None:
                return False
            _copy_run_style(run, suffix_run)

    _set_run_text(fill_run, value)
    fill_run.font.underline = True

    if suffix_run is not None:
        _set_run_text(suffix_run, extra_underscores + suffix)

    return True


def _clear_other_residue(cell, option_labels=('其他', '其它')):
    """清空"其他/其它"选项后面的残留值（模板预填如"电力专网"等示例）。
    user 未选 9/99 时调用，避免残留示例显示。
    """
    for para in cell.paragraphs:
        runs = para.runs
        for i, run in enumerate(runs):
            if not any(lbl in run.text for lbl in option_labels):
                continue
            # 在 label run 自身有 _+ 占位时，不清（_fill_other_option 走这条路）
            if '_' in run.text:
                continue
            for j in range(i + 1, len(runs)):
                tgt = runs[j]
                if tgt._element.find(qn('w:sym')) is not None:
                    break
                if tgt.text.strip() != '':
                    _set_run_text(tgt, '')
            return True
    return False


def _fill_other_option(cell, option_label, text):
    """在"其他_____"这类选项后填充补充文本。

    兼容两种 run 结构：
    - 「其他」和下划线在同一 run（如 '其他______'）→ 直接在该 run 内替换下划线
    - 「其他」后跟独立的空白/下划线 run → 在后续 run 内填值
    填入后继续清空到下一个 sym 之间的残留 run（避免模板预填的"本单位"等示例值被并列保留）。
    """
    if not text:
        return False
    for para in cell.paragraphs:
        runs = para.runs
        for i, run in enumerate(runs):
            if option_label not in run.text:
                continue
            wrote = False
            # keyword run 自带下划线占位
            if '_' in run.text:
                if _fill_placeholder_run(run, text):
                    wrote = True
            for j in range(i + 1, len(runs)):
                target = runs[j]
                # 碰到下一个选项的 sym 就停
                if target._element.find(qn('w:sym')) is not None:
                    break
                if not wrote:
                    if '_' in target.text:
                        if _fill_placeholder_run(target, text):
                            wrote = True
                            continue
                    if target.text.strip() == '':
                        _set_run_text(target, text)
                        wrote = True
                        continue
                # 已经写过 / 找不到合适位置，但仍要清空残留 non-empty run（防止模板示例值保留）
                if wrote and target.text.strip() != '':
                    _set_run_text(target, '')
            return wrote
    return False


def _fill_underline_in_paragraph(paragraph, index, value):
    """将段落中第 index 个下划线 run 替换为值。"""
    count = -1
    for run in paragraph.runs:
        if '_' not in run.text:
            continue
        count += 1
        if count == index:
            _fill_placeholder_run(run, value)
            return True
    return False


def _fill_numbered_lines(cell, values):
    """按段落顺序填充多行（每段一个 value）。
    兼容 3 种模板结构：
      a) 单 run 内 label+下划线（'数据来源单位2___'）→ 替换 _+ 为 value
      b) label run + 后续空白/下划线/残留 run → 首个非 label run 写 value，其余清空
      c) 用户未填该段 → 整段非 label 区也清空
    """
    values = [str(v).strip() if v else '' for v in (values or [])]
    paragraphs = cell.paragraphs
    for idx, para in enumerate(paragraphs):
        value = values[idx] if idx < len(values) else ''
        runs = para.runs
        if not runs:
            continue
        # 跳过段首空 run / sym 编号 run
        i = 0
        while i < len(runs):
            r = runs[i]
            if r.text.strip() == '' or r._element.find(qn('w:sym')) is not None:
                i += 1
                continue
            break
        if i >= len(runs):
            continue
        # 情况 a) label run 自身含 _+ 占位
        if '_' in runs[i].text:
            if value:
                _fill_placeholder_run(runs[i], value, keep_length=True)
            else:
                # 用户未填 → 保留下划线占位（不破坏视觉填空线）
                pass
            # 清空后续非 sym non-empty 残留
            for j in range(i + 1, len(runs)):
                tgt = runs[j]
                if tgt._element.find(qn('w:sym')) is not None:
                    break
                if tgt.text.strip() != '':
                    _set_run_text(tgt, '')
            continue
        # 情况 b)：label 在 i，后续 run 写 value
        wrote = False
        for j in range(i + 1, len(runs)):
            tgt = runs[j]
            if tgt._element.find(qn('w:sym')) is not None:
                break
            if not wrote:
                if '_' in tgt.text and value:
                    _fill_placeholder_run(tgt, value)
                else:
                    _set_run_text(tgt, value)
                wrote = True
            else:
                _set_run_text(tgt, '')
        # 没找到后续位置但有值 → 在 label run 末尾追加（极少见）
        if not wrote and value:
            runs[i].text = runs[i].text + value
    return True


def _fill_option_line(cell, code, value):
    """填充选项行（如 '1 本单位机房 ____'）。兼容下划线占位 / 空格占位 / 河津残留。"""
    code = str(code or "").strip()
    value = str(value or "").strip()
    if not code:
        return False
    for para in cell.paragraphs:
        para_text = "".join(run.text for run in para.runs).strip()
        if not para_text.startswith(code):
            continue
        runs = para.runs
        # 找 label run = 第一个含中文（即非纯空白且非纯数字）run
        label_idx = None
        for i, r in enumerate(runs):
            txt = r.text.strip()
            if not txt or txt == code:
                continue
            label_idx = i
            break
        if label_idx is None:
            continue
        # 情况 a) label run 自身含 _+
        if '_' in runs[label_idx].text and value:
            _fill_placeholder_run(runs[label_idx], value, keep_length=True)
            # 清空后续非 sym 残留
            for j in range(label_idx + 1, len(runs)):
                tgt = runs[j]
                if tgt._element.find(qn('w:sym')) is not None:
                    break
                if tgt.text.strip() != '':
                    _set_run_text(tgt, '')
            return True
        # 情况 b) label 后跟独立 value run
        wrote = False
        for j in range(label_idx + 1, len(runs)):
            tgt = runs[j]
            if tgt._element.find(qn('w:sym')) is not None:
                break
            if not wrote:
                if '_' in tgt.text and value:
                    _fill_placeholder_run(tgt, value)
                else:
                    _set_run_text(tgt, value)
                wrote = True
            else:
                _set_run_text(tgt, '')
        return wrote or not value
    return False


def _safe_set_value(table, row, col, text):
    """
    安全设置「纯值单元格」的文本。
    适用于整个单元格就是值的情况（如单位名称、信用代码等）。
    保留第一个 run 的格式，清除其他 run。
    空值时也会清空 cell 原文本（覆盖模板预填的河津/邵家岭等示例残留）。
    """
    value = "" if text is None else str(text).strip()
    try:
        cell = table.rows[row].cells[col]
        for p in cell.paragraphs:
            if p.runs:
                style_run = p.runs[0]
                _set_run_text(p.runs[0], value)
                _copy_run_style(style_run, p.runs[0])
                _set_run_text(p.runs[0], value)
                for run in p.runs[1:]:
                    run.text = ''
                return
        # 没有 run 则新建
        if cell.paragraphs and value:
            run = cell.paragraphs[0].add_run(value)
            if cell.paragraphs[0].runs:
                _copy_run_style(cell.paragraphs[0].runs[0], run)
            _set_run_text(run, value)
    except (IndexError, AttributeError):
        pass


# ══════════════════════════════════════════════════════
#  备案表生成
# ══════════════════════════════════════════════════════

def generate_beian(template_path: str, output_path: str, data: ProjectData, highlighted_fields=None):
    """基于新备案表模版生成填充后的备案表（精准填充）"""
    if highlighted_fields is None:
        highlighted_fields = []
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    if len(doc.tables) < 7:
        raise ValueError(f"模版表格数不足: 期望>=7, 实际{len(doc.tables)}")

    _set_fill_mode('beian')
    try:
        _fill_beian_internal(doc, data, highlighted_fields)
        doc.save(output_path)
    finally:
        _set_fill_mode('report')
    return output_path


def _fill_beian_internal(doc, data: ProjectData, highlighted_fields):
    """实际填充逻辑（拆出来以便 generate_beian 控制 fill mode）。"""
    u = data.unit
    tgt = data.target
    g = data.grading
    target_name = tgt.name or data.project_name
    unit_name = u.unit_name

    _autofill_attachment_names(data, unit_name, target_name)

    # ── 封面：备案单位填写 ──
    # 模板 L12 runs: ['备',' ','案',' ','单',' ','位：','',' ']
    # 受理那行排除（含"受理"或"盖章"），定位到真正的"备案单位："行
    if unit_name:
        for p in doc.paragraphs:
            full = p.text
            if ('备' in full and '案' in full and '单' in full and '位' in full
                    and '受理' not in full and '盖章' not in full):
                # 找到"位："这一 run 之后的所有 run，把 unit_name 写到下一个 run
                label_run = None
                colon_idx = -1
                for i, run in enumerate(p.runs):
                    if '位' in run.text or '：' in run.text or ':' in run.text:
                        label_run = run
                        if '：' in run.text or ':' in run.text:
                            colon_idx = i
                if colon_idx >= 0 and colon_idx + 1 < len(p.runs):
                    target = p.runs[colon_idx + 1]
                    target.text = unit_name
                    # 清空"位："之后所有其它 run（避免末尾残留多余空格）
                    for run in p.runs[colon_idx + 2:]:
                        run.text = ""
                    if label_run:
                        target.font.name = label_run.font.name
                        target.font.size = label_run.font.size
                        target.font.bold = label_run.font.bold
                        if label_run.font.name:
                            _set_east_asia_font(target, label_run.font.name)
                break

    # ── 表标题替换 （ / ） → （定级对象名） ──
    # 同时处理表五标题："表五（  ）定级对象提交材料情况"（括号内是全角空格）
    if target_name:
        for p in doc.paragraphs:
            txt = p.text
            if '（ / ）' in txt:
                for run in p.runs:
                    if '（ / ）' in run.text:
                        run.text = run.text.replace('（ / ）', f'（{target_name}）')
                    elif '/' in run.text and run.text.strip() in ['/', '/ ']:
                        run.text = run.text.replace('/', target_name)
            elif '表五' in txt and '定级对象提交材料' in txt:
                # 把整段第一个 run 替换，其余清空
                new_text = re.sub(r'表五（\s*）', f'表五（{target_name}）', txt)
                if new_text == txt:
                    new_text = txt.replace('表五（', f'表五（{target_name}', 1)
                if p.runs:
                    p.runs[0].text = new_text
                    for run in p.runs[1:]:
                        run.text = ''

    # ══════ 表2: 单位信息 (index 1, 8列) ══════
    t2 = doc.tables[1]

    # 纯值字段：直接替换整个值区域
    _safe_set_value(t2, 0, 1, u.unit_name)      # 单位名称
    _safe_set_value(t2, 1, 1, u.credit_code)     # 信用代码

    # 地址：精准填空（只替换空格占位 run）
    _fill_address(t2.rows[2].cells[1], u.province, u.city, u.county, u.address)

    _safe_set_value(t2, 3, 1, u.postal_code)     # 邮编
    _safe_set_value(t2, 3, 6, u.admin_code)       # 行政区划代码

    # 负责人
    _safe_set_value(t2, 4, 2, u.leader.name)
    _safe_set_value(t2, 4, 6, u.leader.title)
    _safe_set_value(t2, 5, 2, u.leader.office_phone)
    _safe_set_value(t2, 5, 6, u.leader.email)

    # 安全责任部门
    _safe_set_value(t2, 6, 1, u.security_dept)
    _safe_set_value(t2, 7, 2, u.security_contact.name)
    _safe_set_value(t2, 7, 6, u.security_contact.title)
    _safe_set_value(t2, 8, 2, u.security_contact.office_phone)
    _safe_set_value(t2, 8, 6, u.security_contact.email)
    _safe_set_value(t2, 9, 2, u.security_contact.mobile)
    _safe_set_value(t2, 9, 6, u.security_contact.email)

    # 数据安全部门
    _safe_set_value(t2, 10, 1, u.data_dept)
    _safe_set_value(t2, 11, 2, u.data_contact.name)
    _safe_set_value(t2, 11, 6, u.data_contact.title)
    _safe_set_value(t2, 12, 2, u.data_contact.office_phone)
    _safe_set_value(t2, 12, 6, u.data_contact.email)
    _safe_set_value(t2, 13, 2, u.data_contact.mobile)
    _safe_set_value(t2, 13, 6, u.data_contact.email)

    # 隶属关系（勾选）
    if u.affiliation:
        code = u.affiliation.split('-')[0] if '-' in u.affiliation else u.affiliation
        _find_and_check(t2.rows[14].cells[1], code)

    # 单位类型（勾选）
    if u.unit_type:
        code = u.unit_type.split('-')[0] if '-' in u.unit_type else u.unit_type
        _find_and_check(t2.rows[15].cells[1], code)

    # 行业类别（勾选）- 行16的单元格结构比较特殊，数字编号在不同run中
    if u.industry:
        code = u.industry.split('-')[0] if '-' in u.industry else u.industry
        _find_and_check(t2.rows[16].cells[1], code)

    # 定级对象数量（自动补"个"后缀，未填值则置 "0个"，保持原模板视觉一致）
    def _cnt(v):
        v = str(v or "").strip()
        if not v:
            return "0个"
        return v if v.endswith("个") else f"{v}个"

    _safe_set_value(t2, 17, 1, _cnt(u.current_total))
    _safe_set_value(t2, 17, 3, _cnt(u.current_level2))
    _safe_set_value(t2, 17, 7, _cnt(u.current_level3))
    _safe_set_value(t2, 18, 3, _cnt(u.current_level4))
    _safe_set_value(t2, 18, 7, _cnt(u.current_level5))
    _safe_set_value(t2, 19, 1, _cnt(u.all_total))
    _safe_set_value(t2, 19, 3, _cnt(u.all_level1))
    _safe_set_value(t2, 19, 7, _cnt(u.all_level2))
    _safe_set_value(t2, 20, 3, _cnt(u.all_level3))
    _safe_set_value(t2, 20, 7, _cnt(u.all_level4))
    _safe_set_value(t2, 21, 3, _cnt(u.all_level5))

    # ══════ 表3: 定级对象 (index 2, 9列) ══════
    t3 = doc.tables[2]
    _safe_set_value(t3, 0, 2, tgt.name)

    # 编号填入各独立单元格
    code = tgt.code
    row0 = t3.rows[0].cells
    seen = set()
    code_idx = 0
    for i in range(4, len(row0)):
        cell_id = id(row0[i])
        if cell_id not in seen:
            seen.add(cell_id)
            if code_idx < len(code):
                _safe_set_value(t3, 0, i, code[code_idx])
                code_idx += 1

    # 定级对象类型（勾选）
    type_cell = t3.rows[1].cells[2]
    if tgt.target_type:
        _find_and_check(type_cell, tgt.target_type)
    # 技术类型多选
    if tgt.tech_type:
        for tech in tgt.tech_type.split(','):
            tech = tech.strip()
            if tech:
                _find_and_check(type_cell, tech, multi=True)

    # 业务类型（勾选）
    if tgt.biz_type:
        code = tgt.biz_type.split('-')[0] if '-' in tgt.biz_type else tgt.biz_type
        _find_and_check(t3.rows[2].cells[2], code)
        if code == '9' and tgt.biz_type_other:
            _fill_other_option(t3.rows[2].cells[2], '其他', tgt.biz_type_other)
        else:
            _clear_other_residue(t3.rows[2].cells[2])

    # 业务描述（纯值）
    _safe_set_value(t3, 3, 2, tgt.biz_desc)

    # 服务范围（勾选）
    if tgt.service_scope:
        code = tgt.service_scope.split('-')[0] if '-' in tgt.service_scope else tgt.service_scope
        _find_and_check(t3.rows[4].cells[2], code)
        if code == '99' and tgt.service_scope_other:
            _fill_other_option(t3.rows[4].cells[2], '其它', tgt.service_scope_other)
        else:
            _clear_other_residue(t3.rows[4].cells[2])

    # 服务对象（勾选）
    if tgt.service_target:
        code = tgt.service_target.split('-')[0] if '-' in tgt.service_target else tgt.service_target
        _find_and_check(t3.rows[5].cells[2], code)
        if code == '9' and tgt.service_target_other:
            _fill_other_option(t3.rows[5].cells[2], '其他', tgt.service_target_other)
        else:
            _clear_other_residue(t3.rows[5].cells[2])

    # 部署范围（勾选）
    if tgt.deploy_scope:
        code = tgt.deploy_scope.split('-')[0] if '-' in tgt.deploy_scope else tgt.deploy_scope
        _find_and_check(t3.rows[6].cells[2], code)
        if code == '9' and tgt.deploy_scope_other:
            _fill_other_option(t3.rows[6].cells[2], '其他', tgt.deploy_scope_other)
        else:
            _clear_other_residue(t3.rows[6].cells[2])

    # 网络性质（勾选）
    if tgt.network_type:
        code = tgt.network_type.split('-')[0] if '-' in tgt.network_type else tgt.network_type
        _find_and_check(t3.rows[7].cells[2], code)
        if code == '2':
            _fill_numbered_lines(t3.rows[7].cells[2], [tgt.source_ip, tgt.domain, tgt.protocol_port])
        if code == '9' and tgt.network_type_other:
            _fill_other_option(t3.rows[7].cells[2], '其他', tgt.network_type_other)

    # 网络互联（勾选）
    if tgt.interconnect:
        code = tgt.interconnect.split('-')[0] if '-' in tgt.interconnect else tgt.interconnect
        _find_and_check(t3.rows[8].cells[2], code)
        if code == '9' and tgt.interconnect_other:
            _fill_other_option(t3.rows[8].cells[2], '其它', tgt.interconnect_other)

    # 运行时间（纯值）
    _fill_date_cell(t3.rows[9].cells[2], tgt.run_date)

    # 是否分系统（勾选）
    if tgt.is_subsystem:
        _find_and_check(t3.rows[10].cells[2], tgt.is_subsystem)

    # 上级系统
    _safe_set_value(t3, 11, 2, tgt.parent_system or '/')
    _safe_set_value(t3, 12, 2, tgt.parent_unit or '/')

    # ══════ 表4: 定级等级 (index 3, 5列) ══════
    t4 = doc.tables[3]

    biz_row_map = {"第一级": 1, "第二级": 2, "第三级": 3, "第四级": 4, "第五级": 5}
    svc_row_map = {"第一级": 6, "第二级": 7, "第三级": 8, "第四级": 9, "第五级": 10}

    if g.biz_level in biz_row_map:
        _check_paragraph_items(t4.rows[biz_row_map[g.biz_level]].cells[1], g.biz_level_items)
        _check_first_sym_in_paragraph(t4.rows[biz_row_map[g.biz_level]].cells[4].paragraphs[0], True)
    if g.service_level in svc_row_map:
        _check_paragraph_items(t4.rows[svc_row_map[g.service_level]].cells[1], g.service_level_items)
        _check_first_sym_in_paragraph(t4.rows[svc_row_map[g.service_level]].cells[4].paragraphs[0], True)
    if g.final_level:
        _find_and_check(t4.rows[11].cells[2], g.final_level, multi=True)

    # 定级时间
    _fill_date_cell(t4.rows[12].cells[2], g.grading_date)

    # 定级报告（勾选有/无 + 填附件名）
    report_cell = t4.rows[13].cells[2]
    if g.has_report:
        _find_and_check(report_cell, '有')
    else:
        _find_and_check(report_cell, '无')
    if g.report_name:
        _fill_after_keyword(report_cell, '附件名称', g.report_name)

    # 专家评审（勾选 + 附件名）
    review_cell = t4.rows[14].cells[2]
    if g.has_review:
        _find_and_check(review_cell, '已评审')
    else:
        _find_and_check(review_cell, '未评审')
    if g.review_name:
        _fill_after_keyword(review_cell, '附件名称', g.review_name)

    # 上级主管部门
    supervisor_cell = t4.rows[15].cells[2]
    if g.has_supervisor:
        _find_and_check(supervisor_cell, '有')
    else:
        _find_and_check(supervisor_cell, '无')

    _safe_set_value(t4, 16, 2, g.supervisor_name or '/')
    audit_cell = t4.rows[17].cells[2]
    audit_status = g.supervisor_review_status or ('已审核' if g.supervisor_reviewed else '未审核')
    _find_and_check(audit_cell, audit_status)
    if audit_status == '已审核' and g.supervisor_doc:
        _fill_after_keyword(audit_cell, '附件名称', g.supervisor_doc)

    # 填表人 / 填表日期
    _safe_set_value(t4, 18, 0, f"填表人：{g.filler}")
    _fill_date_cell(t4.rows[18].cells[3], g.fill_date)

    # ══════ 表5: 应用场景 (index 4) — 保持模版默认，仅填已有数据 ══════
    if len(doc.tables) > 4:
        t5 = doc.tables[4]
        sc = data.scenario
        if sc.cloud.enabled:
            _find_and_check(t5.rows[0].cells[2], '是')
            if sc.cloud.role:
                if sc.cloud.role == '二者均勾选':
                    _find_and_check(t5.rows[1].cells[2], '云服务商', multi=True)
                    _find_and_check(t5.rows[1].cells[2], '云服务客户', multi=True)
                else:
                    _find_and_check(t5.rows[1].cells[2], sc.cloud.role, multi=True)
            if sc.cloud.service_model:
                _find_and_check(t5.rows[2].cells[2], sc.cloud.service_model)
                if sc.cloud.service_model == '其他' and sc.cloud.service_model_other:
                    _fill_other_option(t5.rows[2].cells[2], '其他', sc.cloud.service_model_other)
            if sc.cloud.deploy_model:
                _find_and_check(t5.rows[3].cells[2], sc.cloud.deploy_model)
                if sc.cloud.deploy_model == '其他' and sc.cloud.deploy_model_other:
                    _fill_other_option(t5.rows[3].cells[2], '其他', sc.cloud.deploy_model_other)
            if sc.cloud.role in ('云服务商', '二者均勾选'):
                _fill_underline_field(t5.rows[5].cells[2], sc.cloud.provider_scale)
                _safe_set_value(t5, 6, 2, sc.cloud.infra_location)
                _safe_set_value(t5, 7, 2, sc.cloud.ops_location)
            if sc.cloud.role in ('云服务客户', '二者均勾选'):
                _fill_cloud_provider_line(
                    t5.rows[9].cells[2],
                    sc.cloud.provider_name,
                    sc.cloud.platform_level or '三级',
                    sc.cloud.platform_name,
                    _normalize_platform_code(sc.cloud.platform_code),
                )
                _safe_set_value(t5, 10, 2, sc.cloud.client_ops_location)
                if sc.cloud.platform_cert:
                    _fill_after_keyword(t5.rows[11].cells[2], '附件', sc.cloud.platform_cert)
        else:
            _find_and_check(t5.rows[0].cells[2], '否')
        _find_and_check(t5.rows[12].cells[2], '是' if sc.mobile.enabled else '否')
        if sc.mobile.enabled:
            _safe_set_value(t5, 13, 2, sc.mobile.app_name)
            if sc.mobile.wireless:
                _find_and_check(t5.rows[14].cells[2], sc.mobile.wireless, multi=True)
            if sc.mobile.terminal:
                _find_and_check(t5.rows[15].cells[2], sc.mobile.terminal, multi=True)
        _find_and_check(t5.rows[16].cells[2], '是' if sc.iot.enabled else '否')
        if sc.iot.enabled:
            _fill_scenario_options(t5.rows[17].cells[2], sc.iot.perception)
            _fill_scenario_options(t5.rows[18].cells[2], sc.iot.transport)
        _find_and_check(t5.rows[19].cells[2], '是' if sc.ics.enabled else '否')
        if sc.ics.enabled:
            _fill_scenario_options(t5.rows[20].cells[2], sc.ics.function_layer)
            _fill_scenario_options(t5.rows[21].cells[2], sc.ics.composition)
        _find_and_check(t5.rows[22].cells[2], '是' if sc.bigdata.enabled else '否')
        if sc.bigdata.enabled:
            _fill_scenario_options(t5.rows[23].cells[2], sc.bigdata.composition)
            if sc.bigdata.cross_border:
                _find_and_check(t5.rows[24].cells[2], sc.bigdata.cross_border, multi=True)
            comp_tokens = [t.strip() for t in (sc.bigdata.composition or '').replace('，', ',').replace('、', ',').split(',') if t.strip()]
            has_platform = any('大数据平台' in tok for tok in comp_tokens)
            has_client = any(('大数据应用' in tok) or ('大数据资源' in tok) for tok in comp_tokens)
            # r26-r28: 大数据平台填写
            if has_platform:
                if sc.bigdata.platform_scale:
                    _fill_underline_field(t5.rows[26].cells[2], sc.bigdata.platform_scale)
                if sc.bigdata.platform_infra:
                    _safe_set_value(t5, 27, 2, sc.bigdata.platform_infra)
                if sc.bigdata.platform_ops:
                    _safe_set_value(t5, 28, 2, sc.bigdata.platform_ops)
            # r30-r31: 大数据应用、大数据资源填写
            if has_client:
                _fill_cloud_provider_line(
                    t5.rows[30].cells[2],
                    sc.bigdata.platform_provider,
                    sc.bigdata.platform_level or '第三级',
                    sc.bigdata.platform_name,
                    _normalize_platform_code(sc.bigdata.platform_code),
                )
                if sc.bigdata.platform_cert:
                    _fill_after_keyword(t5.rows[31].cells[2], '附件', sc.bigdata.platform_cert)

    # ══════ 表6: 附件清单 (index 5) — 勾选有/无 ══════
    if len(doc.tables) > 5:
        t6 = doc.tables[5]
        att = data.attachment
        items = [
            (0, att.topology), (1, att.org_policy), (2, att.design_plan),
            (3, att.product_list), (4, att.service_list), (5, att.supervisor_doc)
        ]
        for row_idx, item in items:
            cell = t6.rows[row_idx].cells[1]
            if item.has_file:
                _find_and_check(cell, '有')
                if item.file_name:
                    _fill_after_keyword(cell, '附件名称', item.file_name)
            else:
                _find_and_check(cell, '无')

    # ══════ 表7: 数据信息 (index 6) ══════
    if len(doc.tables) > 6:
        t7 = doc.tables[6]
        d = data.data
        _safe_set_value(t7, 0, 1, d.data_name)
        if d.data_level:
            code = d.data_level.split('-')[0] if '-' in d.data_level else d.data_level
            _find_and_check(t7.rows[0].cells[3], code)
        _safe_set_value(t7, 1, 1, d.data_category)
        _safe_set_value(t7, 2, 1, d.data_dept)
        _safe_set_value(t7, 2, 3, d.data_person)
        # 个人信息（勾选）
        if d.personal_info:
            code = d.personal_info.split('-')[0] if '-' in d.personal_info else d.personal_info
            _find_and_check(t7.rows[3].cells[1], code)
        # 数据总量 / 月增长量 — 只填数字到下划线处
        if d.total_size or d.total_size_tb or d.total_size_records:
            _fill_total_size_cell(
                t7.rows[4].cells[1],
                d.total_size_tb if d.total_size_unit == 'TB' else d.total_size,
                d.total_size_unit,
                d.total_size_records,
            )
        if d.monthly_growth or d.monthly_growth_tb:
            _fill_month_growth_cell(
                t7.rows[5].cells[1],
                d.monthly_growth_tb if d.monthly_growth_unit == 'TB' else d.monthly_growth,
                d.monthly_growth_unit,
            )
        if d.data_source:
            for source in d.data_source.split(','):
                code = source.strip().split('-')[0]
                if code:
                    _find_and_check(t7.rows[6].cells[1], code, multi=True)
            if '9-' in d.data_source and d.data_source_other:
                _fill_other_option(t7.rows[6].cells[1], '其他', d.data_source_other)
        _fill_numbered_lines(t7.rows[7].cells[1], [item.strip() for item in d.inflow_units.splitlines() if item.strip()])
        _fill_numbered_lines(t7.rows[8].cells[1], [item.strip() for item in d.outflow_units.splitlines() if item.strip()])
        if d.interaction:
            code = d.interaction.split('-')[0]
            _find_and_check(t7.rows[9].cells[1], code)
            if code != '4' and d.interaction_other:
                _fill_numbered_lines(t7.rows[9].cells[1], [d.interaction_other])
        if d.storage_type:
            code = d.storage_type.split('-')[0]
            _find_and_check(t7.rows[10].cells[1], code)
            # 取真正的位置名：过滤掉前端误传的 select 显示文本（"5-非云计算平台"等）
            raw_name = (d.storage_cloud_name or d.storage_cloud or '').strip()
            if raw_name and (raw_name == d.storage_type or raw_name.startswith(code + '-') or raw_name.startswith(code + ' ')):
                raw_name = ''
            _fill_option_line(t7.rows[10].cells[1], code, raw_name)
        if d.storage_room:
            code = d.storage_room.split('-')[0]
            _find_and_check(t7.rows[11].cells[1], code)
            _fill_option_line(t7.rows[11].cells[1], code, d.storage_room_name)
        if d.storage_region:
            code = d.storage_region.split('-')[0]
            _find_and_check(t7.rows[12].cells[1], code)
            _fill_option_line(t7.rows[12].cells[1], code, d.storage_region_name)

    # ══════ 标黄处理 — 对用户标记的字段在文档中高亮 ══════
    if highlighted_fields:
        # UI field id → (table_index, row, col) 映射
        field_cell_map = {
            'unit_name': (1, 0, 1), 'credit_code': (1, 1, 1),
            'postal_code': (1, 3, 1), 'admin_code': (1, 3, 6),
            'leader_name': (1, 4, 2), 'leader_title': (1, 4, 6),
            'leader_phone': (1, 5, 2), 'leader_email': (1, 5, 6),
            'sec_dept': (1, 6, 1), 'sec_name': (1, 7, 2), 'sec_title': (1, 7, 6),
            'sec_phone': (1, 8, 2), 'sec_email': (1, 8, 6), 'sec_mobile': (1, 9, 2),
            'data_dept': (1, 10, 1), 'data_name': (1, 11, 2), 'data_title': (1, 11, 6),
            'data_phone': (1, 12, 2), 'data_email': (1, 12, 6), 'data_mobile': (1, 13, 2),
            'target_name': (2, 0, 2), 'biz_desc': (2, 3, 2),
            'run_date': (2, 9, 2), 'parent_sys': (2, 11, 2), 'parent_unit': (2, 12, 2),
            'grading_date': (3, 12, 2), 'report_name': (3, 13, 2), 'review_name': (3, 14, 2),
            'data_name_field': (6, 0, 1), 'data_category': (6, 1, 1),
            'data_sec_dept': (6, 2, 1), 'data_sec_person': (6, 2, 3),
            'total_size': (6, 4, 1), 'monthly_growth': (6, 5, 1),
        }
        for field_id in highlighted_fields:
            if field_id in field_cell_map:
                ti, ri, ci = field_cell_map[field_id]
                try:
                    cell = doc.tables[ti].rows[ri].cells[ci]
                    _highlight_cell(cell)
                except (IndexError, AttributeError):
                    pass


def _autofill_attachment_names(data: ProjectData, unit_name: str, target_name: str):
    """按 单位-系统-XXX 规则自动派生附件/报告/评审名（用户已填则保留）。"""
    unit = (unit_name or "").strip()
    system = (target_name or "").strip()
    if not unit or not system:
        return
    prefix = f"{unit}-{system}"

    def _wrap(suffix):
        return f"《{prefix}-{suffix}》"

    g = data.grading
    if g.has_report and not (g.report_name or "").strip():
        g.report_name = _wrap("定级报告")
    if g.has_review and not (g.review_name or "").strip():
        g.review_name = _wrap("专家评审意见表")

    att = data.attachment
    mapping = [
        (att.topology, "网络拓扑结构及说明"),
        (att.org_policy, "系统安全组织机构及管理制度"),
        (att.design_plan, "安全建设整改方案"),
        (att.product_list, "安全产品清单"),
        (att.service_list, "安全服务清单"),
        (att.supervisor_doc, "主管部门定级文件"),
    ]
    for item, suffix in mapping:
        if item.has_file and not (item.file_name or "").strip():
            item.file_name = _wrap(suffix)


def _fill_address(cell, province, city, county, detail):
    """填充地址单元格。兼容空格占位与模板预填残留两种情况。

    模板 run 结构（新模板已预填河津示例）:
      p[0]: ['山西', '省(自治区、直辖市) ', '运城', '地(区、市、州、盟)']
      p[1]: ['河津', '县(区、市、旗)  详细地址', ' ', '樊村镇固镇村邵家岭光伏电站']
    label run 含 '省(' / '地(' / '县(' / '详细地址'，其余视为可覆盖值 run。
    """
    def _is_label(text):
        return any(k in text for k in ('省(', '地(', '县(', '详细地址'))

    def _write_before(runs, label_key, value):
        for i, r in enumerate(runs):
            if label_key in r.text:
                # 向前在「上一个 label 之后」区间内找：首个非 label run 写入，其余清空
                target_idx = None
                for j in range(i - 1, -1, -1):
                    if _is_label(runs[j].text):
                        break  # 不跨越上一个 label
                    if target_idx is None:
                        target_idx = j
                    else:
                        _set_run_text(runs[j], '')
                if target_idx is not None:
                    _set_run_text(runs[target_idx], value or '')
                    return True
                return False
        return False

    def _write_after(runs, label_key, value):
        for i, r in enumerate(runs):
            if label_key in r.text:
                # 向后在「下一个 label 之前」区间内找：首个非 label run 写入，其余清空
                target_idx = None
                for j in range(i + 1, len(runs)):
                    if _is_label(runs[j].text):
                        break
                    if target_idx is None:
                        target_idx = j
                    else:
                        _set_run_text(runs[j], '')
                if target_idx is not None:
                    _set_run_text(runs[target_idx], value or '')
                    return True
                return False
        return False

    try:
        paras = cell.paragraphs
        if len(paras) < 2:
            return
        p0_runs = paras[0].runs
        p1_runs = paras[1].runs

        # 省 / 市
        _write_before(p0_runs, '省(', province)
        _write_before(p0_runs, '地(', city)
        # 县 / 详细地址
        _write_before(p1_runs, '县(', county)
        _write_after(p1_runs, '详细地址', detail)
    except (IndexError, AttributeError):
        pass


def _fill_space_run_before(runs, label, value):
    """在 runs 列表中找到 label 对应 run，将其前面的空格 run 替换为 value"""
    for i, run in enumerate(runs):
        if run.text.strip() == label or run.text.strip().startswith(label):
            # 往前找空格 run
            for j in range(i - 1, -1, -1):
                if runs[j].text.strip() == '':
                    _set_run_text(runs[j], value)
                    return
            break


def _fill_underline_field(cell, value):
    """填充下划线占位字段（如数据总量"___GB"），只替换下划线部分"""
    for para in cell.paragraphs:
        for run in para.runs:
            if '_' in run.text:
                if _fill_placeholder_run(run, value):
                    return
    # fallback: 如果没找到下划线，在第一个run前追加
    if cell.paragraphs and cell.paragraphs[0].runs:
        first = cell.paragraphs[0].runs[0]
        _set_run_text(first, value + first.text.lstrip())


def _fill_amount_unit_run(run, amount, unit):
    if run is None or not amount:
        return False
    if unit == "TB":
        return _fill_placeholder_run(run, str(amount), r"_+(?=TB)")
    return _fill_placeholder_run(run, str(amount), r"_+(?=GB)")


def _fill_total_size_cell(cell, amount, unit, records):
    amount = str(amount or "").strip()
    records = str(records or "").strip()
    unit = (unit or "GB").strip().upper()
    if not cell.paragraphs:
        return False
    if len(cell.paragraphs) >= 1:
        _check_first_sym_in_paragraph(cell.paragraphs[0], bool(amount))
        runs = cell.paragraphs[0].runs
        if len(runs) > 1:
            _fill_amount_unit_run(runs[1], amount, unit)
    if len(cell.paragraphs) >= 2:
        _check_first_sym_in_paragraph(cell.paragraphs[1], bool(records))
        runs = cell.paragraphs[1].runs
        if len(runs) > 1 and records:
            _fill_placeholder_run(runs[1], records)
    return True


def _fill_month_growth_cell(cell, amount, unit):
    amount = str(amount or "").strip()
    unit = (unit or "GB").strip().upper()
    if not amount or not cell.paragraphs or not cell.paragraphs[0].runs:
        return False
    return _fill_amount_unit_run(cell.paragraphs[0].runs[0], amount, unit)


def _fill_cloud_provider_line(cell, provider_name, level, platform_name, platform_code):
    """按下划线占位顺序填入 4 个值（服务商/等级/名称/编号）。

    兼容两种 run 结构：
    - 多 run：每个 _+ 占位独占一个 run（如新模板云 r9 部分场景）
    - 单 run：一个 run 内含多个 _+ 占位（如新模板大数据 r30）
    """
    values = [str(v or "") for v in [provider_name, level, platform_name, platform_code]]
    idx = 0
    placeholders = []  # 记录空值临时占位，最后还原
    while idx < len(values):
        target_run = None
        for para in cell.paragraphs:
            for run in para.runs:
                if re.search(r"_+", run.text or ""):
                    target_run = run
                    break
            if target_run is not None:
                break
        if target_run is None:
            break
        v = values[idx]
        if v:
            _fill_placeholder_run(target_run, v)
        else:
            text = target_run.text
            m = re.search(r"_+", text)
            sentinel = "＿" * (m.end() - m.start())  # 全角下划线占位，避免再被 r"_+" 命中
            target_run.text = text[:m.start()] + sentinel + text[m.end():]
            placeholders.append(target_run)
        idx += 1
    # 还原 sentinel（全角→半角下划线，保留视觉占位）
    for run in placeholders:
        if run.text:
            run.text = run.text.replace("＿", "_")
    return idx > 0


def _normalize_platform_code(platform_code):
    raw = str(platform_code or "").strip()
    if not raw:
        return ""
    parts = [part for part in raw.split("-") if part]
    if len(parts) >= 2:
        return "-".join(parts[:2])
    return raw


def _check_paragraph_items(cell, items):
    targets = {str(item).strip() for item in (items or []) if str(item).strip()}
    if not targets:
        return False
    for para in cell.paragraphs:
        text = "".join(run.text for run in para.runs if run.text).strip()
        if any(target in text for target in targets):
            _check_first_sym_in_paragraph(para, True)
    return True


# ══════════════════════════════════════════════════════
#  定级报告生成
# ══════════════════════════════════════════════════════

def generate_report(template_path: str, output_path: str,
                    report: ReportInfo, project_name: str, highlighted_fields=None):
    """基于新定级报告模版生成填充后的定级报告"""
    if highlighted_fields is None:
        highlighted_fields = []
    _set_fill_mode('report')
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    # ── 第一遍：替换文本、删除说明 ──
    paragraphs_to_remove = []

    # 章节标题 → 该章节正文应替换的字段值
    section_title_to_value = {
        "（一）责任主体": report.responsibility,
        "（二）定级对象构成": report.composition,
        "（三）承载业务": report.business_desc,
        "（四）承载数据": report.carried_data,
        "（五）安全责任": report.security_resp,
        "1、业务信息描述": report.biz_info_desc,
        "2、业务信息受到破坏时所侵害客体的确定": report.biz_victim,
        "3、业务信息受到破坏时对侵害客体的侵害程度的确定": report.biz_degree,
        "1、系统服务描述": report.svc_desc,
        "2、系统服务受到破坏时所侵害客体的确定": report.svc_victim,
        "3、系统服务受到破坏时对侵害客体的侵害程度的确定": report.svc_degree,
    }

    # 按段落顺序划分章节正文（标题段之后、下一个标题之前的所有非空段为该节正文）
    paragraphs_all = list(doc.paragraphs)
    sections_to_apply = []  # [(value, [body_para, ...])]
    current_value = None
    current_body_paras = []

    def _is_other_title(t: str) -> bool:
        return bool(
            re.match(r"^[一二三四五六七八九十]+、", t) or
            re.match(r"^（[一二三四五六七八九十]）", t) or
            re.match(r"^\d+、", t)
        )

    for p in paragraphs_all:
        text = p.text.strip()
        if not text:
            continue
        matched_key = None
        for key in section_title_to_value:
            if text == key or text.startswith(key):
                matched_key = key
                break
        if matched_key is not None:
            if current_value is not None:
                sections_to_apply.append((current_value, current_body_paras))
            current_value = section_title_to_value[matched_key]
            current_body_paras = []
            continue
        # 遇到其它标题，结束当前节
        if _is_other_title(text):
            if current_value is not None:
                sections_to_apply.append((current_value, current_body_paras))
                current_value = None
                current_body_paras = []
            continue
        if current_value is not None:
            current_body_paras.append(p)
    if current_value is not None:
        sections_to_apply.append((current_value, current_body_paras))

    # 应用替换：每节首段替换为新值，其余段标记删除
    for value, body_paras in sections_to_apply:
        if not body_paras:
            continue
        if value:
            _replace_paragraph_text(body_paras[0], value)
            for extra in body_paras[1:]:
                paragraphs_to_remove.append(extra)
        else:
            for extra in body_paras:
                paragraphs_to_remove.append(extra)

    # 标题替换、删除"【填写说明】"、子系统表占位、第X级等
    for p in paragraphs_all:
        text = p.text.strip()
        if not text:
            continue

        # 替换标题（XX → 项目名），保持格式一致
        if "XX" in text and "定级报告" in text:
            for run in p.runs:
                if "XX" in run.text:
                    run.text = run.text.replace("XX", project_name)

        # 清除所有填写说明（【...】标记的内容）
        if text.startswith("【") or "【填写说明" in text or "【描述参考示例】" in text or "【网络边界描述示例】" in text:
            paragraphs_to_remove.append(p)
            continue
        # 清除以"说明："开头的提示段落
        if text.startswith("说明："):
            paragraphs_to_remove.append(p)
            continue
        # 清除模板中残留的"该定级对象是否采用了新技术"等说明
        if "该定级对象是否采用了新技术" in text:
            paragraphs_to_remove.append(p)
            continue

        # 替换子系统表描述
        if "该定级对象包括以下子系统" in text:
            if not report.subsystems:
                paragraphs_to_remove.append(p)
            else:
                _replace_paragraph_text(p, "该定级对象包括以下子系统：")

        # 替换等级占位
        if "第X级" in text:
            new_text = text
            if "最终确定" in text:
                new_text = text.replace("第X级", report.final_level)
                new_text = new_text.replace("XXX", project_name)
            elif "业务信息安全保护等级" in text and "系统服务" not in text:
                new_text = text.replace("第X级", report.biz_level)
            elif "系统服务安全保护等级" in text:
                new_text = text.replace("第X级", report.svc_level)
            else:
                new_text = text.replace("第X级", report.final_level)
            if new_text != text:
                _replace_paragraph_text(p, new_text)
                continue

        # 替换 XXX
        if "XXX" in text and "第X级" not in text:
            _replace_paragraph_text(p, text.replace("XXX", project_name))

    # 删除标记的段落
    for p in paragraphs_to_remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)

    # ── 清除多余空行（连续2个以上空段落压缩为1个） ──
    _remove_consecutive_blanks(doc)

    # ── 插入网络拓扑图（放在"定级对象构成"描述文字之后） ──
    if report.topology_image and os.path.exists(report.topology_image):
        _insert_topology_image(doc, report.topology_image)

    # ── 子系统表格处理 ──
    for table in doc.tables:
        if _cell_text_safe(table, 0, 0) == "序号":
            if report.subsystems:
                while len(table.rows) > 1:
                    table._tbl.remove(table.rows[-1]._tr)
                for sub in report.subsystems:
                    row = table.add_row()
                    set_cell_text(row.cells[0], sub.index)
                    set_cell_text(row.cells[1], sub.name)
                    set_cell_text(row.cells[2], sub.description)
            else:
                # 无子系统时删除表格
                table._tbl.getparent().remove(table._tbl)
            break

    # ── 矩阵表涂色（业务信息 & 系统服务安全保护等级矩阵表） ──
    _shade_matrix_tables(doc, report)

    # ── 最终等级汇总表 ──
    for table in doc.tables:
        if _cell_text_safe(table, 0, 0) == "定级对象名称":
            if len(table.rows) > 1:
                _safe_set_value(table, 1, 0, project_name)
                _safe_set_value(table, 1, 1, report.final_level)
                _safe_set_value(table, 1, 2, report.biz_level)
                _safe_set_value(table, 1, 3, report.svc_level)
            break

    doc.save(output_path)
    return output_path


def _remove_consecutive_blanks(doc):
    """压缩连续空段落，最多保留1个"""
    prev_blank = False
    to_remove = []
    for p in doc.paragraphs:
        is_blank = not p.text.strip()
        if is_blank and prev_blank:
            to_remove.append(p)
        prev_blank = is_blank
    for p in to_remove:
        parent = p._element.getparent()
        if parent is not None:
            parent.remove(p._element)


def _shade_matrix_tables(doc, report):
    """
    对业务信息和系统服务安全保护等级矩阵表进行涂色。
    根据等级在对应行列交叉处涂黑（深色背景+白色文字）。
    """
    matrix_sources = [
        ("业务信息安全被破坏时所侵害的客体", report.biz_victim, report.biz_degree),
        ("系统服务安全被破坏时所侵害的客体", report.svc_victim, report.svc_degree),
    ]
    for header_text, victim_text, degree_text in matrix_sources:
        table = next((item for item in doc.tables if header_text in _cell_text_safe(item, 0, 0)), None)
        if table is None:
            continue
        row_idx = _match_matrix_row(victim_text)
        col_idx = _match_matrix_col(degree_text)
        if row_idx is None or col_idx is None:
            continue
        _shade_cell_gray(table.rows[row_idx].cells[0])
        _shade_cell_gray(table.rows[1].cells[col_idx])
        _shade_cell_gray(table.rows[row_idx].cells[col_idx])


def _match_matrix_row(text):
    raw = str(text or "")
    if "公民" in raw or "法人" in raw or "其他组织" in raw:
        return 2
    if "社会秩序" in raw or "公共利益" in raw:
        return 3
    if "国家安全" in raw or "地区安全" in raw or "国计民生" in raw:
        return 4
    return None


def _match_matrix_col(text):
    raw = str(text or "")
    if "特别严重" in raw:
        return 3
    if "严重" in raw:
        return 2
    if "一般" in raw:
        return 1
    return None


def _shade_cell_gray(cell):
    """给单元格设置灰色背景、白色文字。"""
    # 设置单元格底纹
    tc_pr = cell._element.get_or_add_tcPr()
    for node in list(tc_pr):
        if node.tag == qn('w:shd'):
            tc_pr.remove(node)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), '808080')
    tc_pr.append(shading)
    # 设置文字为白色
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _replace_paragraph_text(paragraph, new_text):
    """替换段落文本，保留第一个 run 的格式。
    若 new_text 含换行（用户在 textarea 按回车），按行拆分为多段：
    第一行写入 paragraph，后续行在它之后插入新段落，继承原段落 pPr（含首行缩进）。
    """
    placeholder_pattern = re.compile(r"(身份证号码[:：][X_]{3,}|X{4,}|_{3,})")
    style_run = paragraph.runs[0] if paragraph.runs else None
    if not paragraph.runs:
        paragraph.add_run("")
        style_run = paragraph.runs[0]

    # 拆行：兼容 \r\n / \n / \r
    raw = (new_text or "")
    lines = re.split(r"\r\n|\r|\n", raw)
    if not lines:
        lines = [""]

    def _write_line_to_paragraph(para, text):
        for run in para.runs:
            run.text = ""
        parts = [part for part in placeholder_pattern.split(text) if part]
        if not parts:
            parts = [text]
        for idx, part in enumerate(parts):
            run = para.runs[0] if (idx == 0 and para.runs) else para.add_run()
            run.text = part
            if style_run is not None:
                _copy_run_style(style_run, run)
            if placeholder_pattern.fullmatch(part):
                _highlight_run(run, "yellow")

    # 第一行写入原段落
    _write_line_to_paragraph(paragraph, lines[0])

    # 后续行在原段落后插入新段落（继承 pPr）
    src_pPr = paragraph._element.find(qn('w:pPr'))
    anchor = paragraph._element
    from docx.text.paragraph import Paragraph as _Paragraph
    for line in lines[1:]:
        new_p_el = OxmlElement('w:p')
        if src_pPr is not None:
            new_p_el.insert(0, deepcopy(src_pPr))
        anchor.addnext(new_p_el)
        anchor = new_p_el
        new_para = _Paragraph(new_p_el, paragraph._parent)
        new_para.add_run("")
        _write_line_to_paragraph(new_para, line)


def _cell_text_safe(table, row, col):
    try:
        return table.rows[row].cells[col].text.strip()
    except (IndexError, AttributeError):
        return ""


def _insert_topology_image(doc, image_path):
    """在定级对象构成章节的描述文字之后插入拓扑图。
    若目标段已含模板预填的占位图（如示例拓扑图），先清空再插入用户图，避免重复。"""
    def _clear_drawings(para):
        """删除段内所有 w:drawing 元素（连带其所属 run），保留 pPr。"""
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        for r in list(para._element.findall(f'{ns}r')):
            if r.find(f'{ns}drawing') is not None or r.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip') is not None:
                para._element.remove(r)

    found_section = False
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if ("定级对象构成" in text and "（二）" in text) or "网络拓扑图" in text:
            found_section = True
            continue
        if found_section:
            # 跳过描述性文字段落，找到第一个空段落或下一个标题之前的位置
            if not text or text.startswith("（三）") or text.startswith("（四）"):
                target_p = doc.paragraphs[i]
                if not text:
                    _clear_drawings(target_p)
                    run = target_p.add_run()
                    run.add_picture(image_path, width=Inches(5.5))
                else:
                    para = p.insert_paragraph_before()
                    run = para.add_run()
                    run.add_picture(image_path, width=Inches(5.5))
                return

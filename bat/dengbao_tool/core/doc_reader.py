"""旧文档数据提取 + 勾选项读取 + 调查表读取"""

import os
import tempfile
import re
from docx import Document
from docx.oxml.ns import qn
from models.project_data import (
    ProjectData, UnitInfo, ContactInfo, TargetInfo, GradingInfo,
    ReportInfo, SubSystem, ScenarioInfo, CloudInfo, DataInfo,
    AttachmentInfo, AttachmentItem
)


def _cell_text(table, row, col):
    """安全获取表格单元格文本"""
    try:
        return table.rows[row].cells[col].text.strip()
    except (IndexError, AttributeError):
        return ""


def _read_checked_options(cell):
    """
    读取单元格中所有被勾选(☑)的选项。
    返回列表: [(编号, 文字), ...]
    勾选标记: w:sym font=Wingdings char=F0FE
    """
    results = []
    for para in cell.paragraphs:
        runs = para.runs
        for i, run in enumerate(runs):
            sym = run._element.find(qn('w:sym'))
            if sym is None:
                continue
            char = sym.get(qn('w:char'))
            # 勾选标记：F0FE 或 00FE (不同版本 Word 可能不同)
            if char not in ('F0FE', '00FE'):
                continue
            # 取 sym 后面的选项编号和文字
            parts = []
            for j in range(i + 1, min(i + 4, len(runs))):
                if runs[j]._element.find(qn('w:sym')) is not None:
                    break
                t = runs[j].text.strip()
                if t:
                    parts.append(t)
            if parts:
                results.append(parts)
    return results


def _read_first_checked(cell):
    """读取单元格中第一个被勾选的选项，返回编号字符串"""
    opts = _read_checked_options(cell)
    if opts:
        return opts[0][0]  # 返回编号
    return ""


def _parse_address_runs(cell):
    """从地址单元格的 run 结构解析省、市、县、详细地址"""
    province = city = county = detail = ""
    try:
        paras = cell.paragraphs
        if len(paras) >= 2:
            p0 = paras[0].runs
            p1 = paras[1].runs
            # 省名在"省"字前的 run 中
            for i, run in enumerate(p0):
                if run.text.strip() == '省' or run.text.strip().startswith('省'):
                    # 往前找有内容的 run
                    for j in range(i - 1, -1, -1):
                        t = p0[j].text.strip()
                        if t and t not in ['', ' ']:
                            province = t
                            break
                    break
            # 市名在"地"字前
            for i, run in enumerate(p0):
                if run.text.strip() == '地' or run.text.strip().startswith('地'):
                    for j in range(i - 1, -1, -1):
                        t = p0[j].text.strip()
                        if t and t not in ['', ' ', ')'] and '自治区' not in t:
                            city = t
                            break
                    break
            # 县名在"县"字前
            for i, run in enumerate(p1):
                if run.text.strip() == '县' or run.text.strip().startswith('县'):
                    for j in range(i - 1, -1, -1):
                        t = p1[j].text.strip()
                        if t and t not in ['', ' ']:
                            county = t
                            break
                    break
            # 详细地址在"详细地址"后，可能跨多个run
            for i, run in enumerate(p1):
                if '详细地址' in run.text:
                    parts = []
                    for j in range(i + 1, len(p1)):
                        t = p1[j].text.strip()
                        if t:
                            parts.append(t)
                    if parts:
                        detail = ''.join(parts)
                    break
    except (IndexError, AttributeError):
        pass
    return province, city, county, detail


def _parse_cloud_provider_line(text: str):
    """解析“云服务商为___ 平台安全等级___ 平台名称___ 平台备案编号___”结构。"""
    raw = (text or "").replace("\n", " ").strip()
    result = {
        "provider_name": "",
        "platform_level": "",
        "platform_name": "",
        "platform_code": "",
    }
    if not raw:
        return result

    patterns = {
        "provider_name": r"云服务商为\s*(.*?)\s*平台安全等级",
        "platform_level": r"平台安全等级\s*(.*?)\s*平台名称",
        "platform_name": r"平台名称\s*(.*?)\s*平台备案编号",
        "platform_code": r"平台备案编号\s*([A-Za-z0-9\-]+)",
    }
    for key, pattern in patterns.items():
        match = re.search(pattern, raw)
        if match:
            result[key] = match.group(1).strip(" _")

    if not result["provider_name"] and "云服务商为" not in raw:
        result["provider_name"] = raw
    return result


def _clean_inline_token(text: str, prefix: str = "") -> str:
    raw = re.sub(r"\s+", "", (text or ""))
    if prefix and raw.startswith(prefix):
        raw = raw[len(prefix):]
    return raw.strip(" _")


def _parse_dual_unit_amount(text: str, prefix: str = ""):
    """
    解析形如:
    - 1_____GB/_____TB
    - 130GB/_____TB
    - 1_____GB/8TB
    - ___________GB/2TB
    返回 (amount, unit)
    """
    raw = re.sub(r"\s+", "", (text or ""))
    match = re.search(r"^(.*?)GB/(.*?)TB", raw)
    if not match:
        cleaned = _clean_inline_token(raw, prefix)
        return cleaned, "GB" if cleaned else ""

    gb_part = match.group(1)
    tb_part = match.group(2)
    if prefix and gb_part.startswith(prefix):
        gb_part = gb_part[len(prefix):]

    gb_value = gb_part.strip(" _")
    tb_value = tb_part.strip(" _")
    if gb_value:
        return gb_value, "GB"
    if tb_value:
        return tb_value, "TB"
    return "", ""


def _parse_records_amount(text: str, prefix: str = "2") -> str:
    raw = re.sub(r"\s+", "", (text or ""))
    match = re.search(rf"^{re.escape(prefix)}?(.*?)万条", raw)
    if not match:
        return _clean_inline_token(raw, prefix)
    return match.group(1).strip(" _")


def read_beian_docx(path: str) -> ProjectData:
    """
    从备案表 .docx 提取数据（含勾选项解析）。
    """
    doc = Document(path)
    data = ProjectData()

    if len(doc.tables) < 3:
        return data

    # === 表2: 单位信息 (index 1, 8列) ===
    t2 = doc.tables[1]
    u = data.unit
    u.unit_name = _cell_text(t2, 0, 1)
    u.credit_code = _cell_text(t2, 1, 1)

    # 地址：精准解析省市县
    u.province, u.city, u.county, u.address = _parse_address_runs(t2.rows[2].cells[1])

    u.postal_code = _cell_text(t2, 3, 1)
    u.admin_code = _cell_text(t2, 3, 6)

    # 单位负责人
    u.leader.name = _cell_text(t2, 4, 2)
    u.leader.title = _cell_text(t2, 4, 6)
    u.leader.office_phone = _cell_text(t2, 5, 2)
    u.leader.email = _cell_text(t2, 5, 6)

    # 安全责任部门
    u.security_dept = _cell_text(t2, 6, 1)
    u.security_contact.name = _cell_text(t2, 7, 2)
    u.security_contact.title = _cell_text(t2, 7, 6)
    u.security_contact.office_phone = _cell_text(t2, 8, 2)
    u.security_contact.email = _cell_text(t2, 8, 6)
    u.security_contact.mobile = _cell_text(t2, 9, 2)

    # 数据安全部门
    if len(t2.rows) > 13:
        u.data_dept = _cell_text(t2, 10, 1)
        u.data_contact.name = _cell_text(t2, 11, 2)
        u.data_contact.title = _cell_text(t2, 11, 6)
        u.data_contact.office_phone = _cell_text(t2, 12, 2)
        u.data_contact.email = _cell_text(t2, 12, 6)
        u.data_contact.mobile = _cell_text(t2, 13, 2)

    # 勾选项读取
    if len(t2.rows) > 16:
        u.affiliation = _read_first_checked(t2.rows[14].cells[1])
        u.unit_type = _read_first_checked(t2.rows[15].cells[1])
        u.industry = _read_first_checked(t2.rows[16].cells[1])

    # 定级对象数量
    if len(t2.rows) > 21:
        u.current_total = _cell_text(t2, 17, 1)
        u.current_level2 = _cell_text(t2, 17, 3)
        u.current_level3 = _cell_text(t2, 17, 7)
        u.current_level4 = _cell_text(t2, 18, 3)
        u.current_level5 = _cell_text(t2, 18, 7)
        u.all_total = _cell_text(t2, 19, 1)
        u.all_level1 = _cell_text(t2, 19, 3)
        u.all_level2 = _cell_text(t2, 19, 7)
        u.all_level3 = _cell_text(t2, 20, 3)
        u.all_level4 = _cell_text(t2, 20, 7)
        u.all_level5 = _cell_text(t2, 21, 3)

    # === 表3: 定级对象信息 (index 2, 9列) ===
    t3 = doc.tables[2]
    tgt = data.target
    tgt.name = _cell_text(t3, 0, 2)

    # 编号
    code_parts = []
    row0 = t3.rows[0].cells
    seen = set()
    for i in range(4, len(row0)):
        t = row0[i].text.strip()
        cell_id = id(row0[i])
        if cell_id not in seen and t:
            seen.add(cell_id)
            code_parts.append(t)
    tgt.code = "".join(code_parts)

    # 勾选项
    if len(t3.rows) > 1:
        # 对象类型在 cells[2]（cells[0-1] 是合并的标签）
        type_opts = _read_checked_options(t3.rows[1].cells[2])
        if type_opts:
            main_types = []
            tech_types = []
            for parts in type_opts:
                joined = ''.join(parts)
                if '信息系统' in joined or '通信网络' in joined or '数据资源' in joined:
                    main_types.append(joined)
                else:
                    tech_types.append(joined)
            if main_types:
                tgt.target_type = main_types[0]
            if tech_types:
                tgt.tech_type = ','.join(tech_types)

    if len(t3.rows) > 2:
        tgt.biz_type = _read_first_checked(t3.rows[2].cells[2])
    if len(t3.rows) > 4:
        tgt.service_scope = _read_first_checked(t3.rows[4].cells[2])
    if len(t3.rows) > 5:
        tgt.service_target = _read_first_checked(t3.rows[5].cells[2])
    if len(t3.rows) > 6:
        tgt.deploy_scope = _read_first_checked(t3.rows[6].cells[2])
    if len(t3.rows) > 7:
        tgt.network_type = _read_first_checked(t3.rows[7].cells[2])
    if len(t3.rows) > 8:
        interconnect = _read_first_checked(t3.rows[8].cells[2])
        if interconnect:
            tgt.interconnect = interconnect
    if len(t3.rows) > 10:
        tgt.is_subsystem = _read_first_checked(t3.rows[10].cells[2])

    tgt.biz_desc = _cell_text(t3, 3, 2)
    tgt.run_date = _cell_text(t3, 9, 2)
    tgt.parent_system = _cell_text(t3, 11, 2)
    tgt.parent_unit = _cell_text(t3, 12, 2)

    # === 表4: 定级等级 (index 3, 5列) ===
    if len(doc.tables) > 3:
        t4 = doc.tables[3]
        g = data.grading
        if len(t4.rows) > 18:
            # 安全等级勾选
            level = _read_first_checked(t4.rows[11].cells[2])
            if level:
                g.final_level = level
                g.biz_level = level
                g.service_level = level

            g.grading_date = _cell_text(t4, 12, 2)

            # 定级报告有/无
            report_opts = _read_checked_options(t4.rows[13].cells[2])
            for parts in report_opts:
                if '有' in parts:
                    g.has_report = True
                elif '无' in parts:
                    g.has_report = False

            # 提取附件名称（在"附件名称"后面）
            report_text = _cell_text(t4, 13, 2)
            if '附件名称' in report_text:
                name_part = report_text.split('附件名称')[-1].strip()
                if name_part:
                    g.report_name = name_part

            # 专家评审
            review_opts = _read_checked_options(t4.rows[14].cells[2])
            for parts in review_opts:
                if '已评审' in parts:
                    g.has_review = True
                elif '未评审' in parts:
                    g.has_review = False
            review_text = _cell_text(t4, 14, 2)
            if '附件名称' in review_text:
                name_part = review_text.split('附件名称')[-1].strip()
                if name_part:
                    g.review_name = name_part

            supervisor_opts = _read_checked_options(t4.rows[15].cells[2])
            for parts in supervisor_opts:
                joined = ''.join(parts)
                if '有' in joined:
                    g.has_supervisor = True
                if '无' in joined:
                    g.has_supervisor = False

            supervisor_review_opts = _read_checked_options(t4.rows[17].cells[2])
            for parts in supervisor_review_opts:
                joined = ''.join(parts)
                if '已审核' in joined:
                    g.supervisor_reviewed = True
                    g.supervisor_review_status = '已审核'
                elif '未审核' in joined:
                    g.supervisor_reviewed = False
                    g.supervisor_review_status = '未审核'
            supervisor_text = _cell_text(t4, 17, 2)
            if '附件名称' in supervisor_text:
                name_part = supervisor_text.split('附件名称')[-1].strip()
                if name_part:
                    g.supervisor_doc = name_part

            g.filler = _cell_text(t4, 18, 0).replace("填表人：", "")
            g.fill_date = _cell_text(t4, 18, 3).replace("填表日期：", "")

    # === 表5: 应用场景 (index 4) ===
    if len(doc.tables) > 4:
        t5 = doc.tables[4]
        sc = data.scenario
        role_opts = _read_checked_options(t5.rows[1].cells[2])
        role_labels = []
        for parts in role_opts:
            joined = ''.join(parts)
            if '云服务商' in joined:
                role_labels.append('云服务商')
            elif '云服务客户' in joined:
                role_labels.append('云服务客户')
        if len(role_labels) >= 2:
            sc.cloud.role = '二者均勾选'
        elif role_labels:
            sc.cloud.role = role_labels[0]

        cloud_text = _cell_text(t5, 9, 2)
        if cloud_text and cloud_text.strip():
            sc.cloud.enabled = True
            parsed = _parse_cloud_provider_line(cloud_text)
            sc.cloud.provider_name = parsed["provider_name"]
            sc.cloud.platform_level = parsed["platform_level"]
            sc.cloud.platform_name = parsed["platform_name"]
            sc.cloud.platform_code = parsed["platform_code"]
        client_ops = _cell_text(t5, 10, 2)
        if client_ops:
            sc.cloud.client_ops_location = client_ops
        provider_scale = _cell_text(t5, 5, 2)
        if provider_scale:
            sc.cloud.provider_scale = provider_scale.replace('云服务客户数量', '').replace('个', '').strip(' _')
        infra_location = _cell_text(t5, 6, 2)
        if infra_location:
            sc.cloud.infra_location = infra_location
        ops_location = _cell_text(t5, 7, 2)
        if ops_location:
            sc.cloud.ops_location = ops_location
        platform_cert = _cell_text(t5, 11, 2)
        if platform_cert:
            sc.cloud.platform_cert = platform_cert.replace('附件', '').strip(' _')

    # === 表6: 附件清单 (index 5) ===
    if len(doc.tables) > 5:
        t6 = doc.tables[5]
        att = data.attachment
        items = [
            (0, att.topology), (1, att.org_policy), (2, att.design_plan),
            (3, att.product_list), (4, att.service_list), (5, att.supervisor_doc)
        ]
        for row_idx, item in items:
            cell = t6.rows[row_idx].cells[1]
            checked = _read_checked_options(cell)
            for parts in checked:
                if '有' in parts:
                    item.has_file = True
                elif '无' in parts:
                    item.has_file = False
            # 附件名称
            text = _cell_text(t6, row_idx, 1)
            if '附件名称' in text:
                name_part = text.split('附件名称')[-1].strip()
                if name_part:
                    item.file_name = name_part

    # === 表7: 数据信息 (index 6) ===
    if len(doc.tables) > 6:
        t7 = doc.tables[6]
        d = data.data
        d.data_name = _cell_text(t7, 0, 1)
        d.data_category = _cell_text(t7, 1, 1)
        d.data_dept = _cell_text(t7, 2, 1)
        d.data_person = _cell_text(t7, 2, 3)
        # 个人信息勾选
        if len(t7.rows) > 3:
            d.personal_info = _read_first_checked(t7.rows[3].cells[1])
        if len(t7.rows) > 4:
            total_cell = t7.rows[4].cells[1]
            if total_cell.paragraphs:
                total_amount, total_unit = _parse_dual_unit_amount(total_cell.paragraphs[0].text, prefix="1")
                if total_unit == "TB":
                    d.total_size_unit = "TB"
                    d.total_size_tb = total_amount
                    d.total_size = ""
                else:
                    d.total_size_unit = total_unit or "GB"
                    d.total_size = total_amount
                    d.total_size_tb = ""
            if len(total_cell.paragraphs) > 1:
                d.total_size_records = _parse_records_amount(total_cell.paragraphs[1].text, prefix="2")
        if len(t7.rows) > 5:
            month_cell = t7.rows[5].cells[1]
            if month_cell.paragraphs:
                month_amount, month_unit = _parse_dual_unit_amount(month_cell.paragraphs[0].text)
                if month_unit == "TB":
                    d.monthly_growth_unit = "TB"
                    d.monthly_growth_tb = month_amount
                    d.monthly_growth = ""
                else:
                    d.monthly_growth_unit = month_unit or "GB"
                    d.monthly_growth = month_amount
                    d.monthly_growth_tb = ""

    data.project_name = tgt.name or u.unit_name
    return data


def read_report_docx(path: str) -> ReportInfo:
    """从定级报告 .docx 提取数据"""
    doc = Document(path)
    report = ReportInfo()
    paragraphs = doc.paragraphs

    for p in paragraphs:
        if p.text.strip():
            report.system_name = p.text.strip()
            break

    current_section = ""
    section_texts = {}

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue

        if text.startswith("（一）") and ("责任主体" in text or "安全责任" in text):
            current_section = "responsibility"
            continue
        elif text.startswith("（二）") and ("基本要素" in text or "网络结构" in text or "定级对象构成" in text):
            current_section = "composition"
            continue
        elif text.startswith("（三）") and ("业务" in text or "承载" in text):
            current_section = "business"
            continue
        elif text.startswith("（四）") and "安全责任" in text:
            current_section = "security"
            continue
        elif text.startswith("一、") or text.startswith("二、") or text.startswith("三、"):
            current_section = ""
            continue
        elif "1、业务信息描述" in text:
            current_section = "biz_info"
            continue
        elif "2、业务信息受到破坏时所侵害客体" in text:
            current_section = "biz_victim"
            continue
        elif "3、" in text and "业务信息" in text and "程度" in text:
            current_section = "biz_degree"
            continue
        elif "4、业务信息安全" in text and "确定" in text:
            current_section = "biz_level_result"
            continue
        elif "1、系统服务描述" in text:
            current_section = "svc_desc"
            continue
        elif "2、系统服务受到破坏时所侵害客体" in text:
            current_section = "svc_victim"
            continue
        elif "3、" in text and "系统服务" in text and "程度" in text:
            current_section = "svc_degree"
            continue
        elif "4、系统服务安全" in text and "确定" in text:
            current_section = "svc_level_result"
            continue
        elif text.startswith(("1、", "2、", "3、", "4、", "（一）", "（二）", "（三）")):
            current_section = ""
            continue

        if current_section:
            if current_section not in section_texts:
                section_texts[current_section] = []
            section_texts[current_section].append(text)

    report.responsibility = "\n".join(section_texts.get("responsibility", []))
    report.composition = "\n".join(section_texts.get("composition", []))
    report.business_desc = "\n".join(section_texts.get("business", []))
    report.security_resp = "\n".join(section_texts.get("security", []))
    report.biz_info_desc = "\n".join(section_texts.get("biz_info", []))
    report.biz_victim = "\n".join(section_texts.get("biz_victim", []))
    report.biz_degree = "\n".join(section_texts.get("biz_degree", []))
    report.svc_desc = "\n".join(section_texts.get("svc_desc", []))
    report.svc_victim = "\n".join(section_texts.get("svc_victim", []))
    report.svc_degree = "\n".join(section_texts.get("svc_degree", []))

    for p in paragraphs:
        text = p.text.strip()
        if "最终确定" in text and "等级为" in text:
            for level in ["第五级", "第四级", "第三级", "第二级", "第一级"]:
                if level in text:
                    report.final_level = level
                    break
        if "业务信息安全保护等级为" in text or "业务信息安全等级为" in text:
            for level in ["第五级", "第四级", "第三级", "第二级", "第一级"]:
                if level in text:
                    report.biz_level = level
                    break
        if "系统服务安全" in text and "等级为" in text:
            for level in ["第五级", "第四级", "第三级", "第二级", "第一级"]:
                if level in text:
                    report.svc_level = level
                    break

    for table in doc.tables:
        header = _cell_text(table, 0, 0)
        if "序号" in header:
            for i in range(1, len(table.rows)):
                sub = SubSystem(
                    index=_cell_text(table, i, 0),
                    name=_cell_text(table, i, 1),
                    description=_cell_text(table, i, 2)
                )
                if sub.name:
                    report.subsystems.append(sub)
            break

    return report


def read_survey_docx(path: str):
    """
    从等级保护对象基本情况调查表提取拓扑图和描述。
    返回 (image_path, description_text) 或 (None, None)
    """
    if not path or not os.path.exists(path):
        return None, None

    doc = Document(path)
    image_path = None
    desc_text = ""

    # 查找拓扑图图片（在"网络拓扑结构图"段落附近）
    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            drawings = run._element.findall(qn('w:drawing'))
            if drawings:
                # 提取图片
                for d in drawings:
                    blips = d.findall(
                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for b in blips:
                        rId = b.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId:
                            try:
                                rel = doc.part.rels[rId]
                                img_part = rel.target_part
                                ext = '.png' if 'png' in img_part.content_type else '.jpg'
                                tmp = os.path.join(
                                    tempfile.gettempdir(), f'dengbao_topo{ext}')
                                with open(tmp, 'wb') as f:
                                    f.write(img_part.blob)
                                image_path = tmp
                            except Exception:
                                pass

    # 查找拓扑描述文字（"如图2.1所示"开头的段落）
    for p in doc.paragraphs:
        text = p.text.strip()
        if '如图' in text and '所示' in text:
            desc_text = text
            break

    return image_path, desc_text

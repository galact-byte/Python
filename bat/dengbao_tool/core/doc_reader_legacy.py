"""旧版备案表的通用标签扫描读取器。

设计要点：
- 不依赖具体的表索引或行列数，按标签关键字定位字段。
- 同 row 内有合并单元格，把相邻重复 cell 文本去重。
- 嵌套字段（负责人/联系人内部的姓名/电话/邮箱）按"段上下文"prefix 防止 key 冲突。
- 勾选状态用 w:sym (F0FE=Wingdings) 解析。
"""

import re
from typing import Dict, List, Tuple
from docx.oxml.ns import qn
from models.project_data import (
    ProjectData, ContactInfo, ReportInfo, SubSystem,
)


# 顶层标签：作为 section 切换依据
_SECTION_LABELS = {
    "单位\n负责人": "leader",
    "单位负责人": "leader",
    "责任部门\n联系人": "sec_contact",
    "责任部门联系人": "sec_contact",
    "网络安全管理\n机构联系人": "sec_contact",
    "数据安全管理\n部门联系人": "data_contact",
    "信息系统\n总数": "level_count",
    "确定 业务 信息 安全 保护 等级": "biz_level_table",
    "确定 系统 服务 安全 保护 等级": "svc_level_table",
}

# 行内字段标签：在同一行中后接值
_INLINE_LABELS = {
    # 单位维度
    "单位名称", "统一社会信用代码", "信用代码",
    "单位地址", "邮政编码", "行政区划代码",
    "姓 名", "姓名", "职务/职称", "职务", "职称",
    "办公电话", "电子邮件", "移动电话",
    "责任部门", "网络安全管理机构", "数据安全管理部门", "负责人",
    "隶属关系", "单位类型", "行业类别",
    # 定级对象维度
    "系统名称", "定级对象名称", "系统编号", "定级对象编号",
    "业务类型", "业务描述",
    "服务范围", "服务对象", "覆盖范围", "部署范围",
    "网络性质", "系统互联情况", "网络互联情况",
    "运行时间", "投入运行时间", "是否分系统", "上级系统名称", "上级系统所属单位",
    # 等级与流程
    "信息系统安全保护等级", "等级保护对象安全保护等级", "定级时间",
    "业务信息安全保护等级", "系统服务安全保护等级", "最终安全保护等级",
    "定级报告", "系统定级报告", "专家评审情况", "评审情况",
    "是否有主管部门", "主管部门名称", "主管部门审批定级情况", "审核情况",
    "填表人", "填表日期",
    # 数据
    "数据名称", "数据类别", "数据级别", "数据安全责任部门", "数据安全负责人",
}


def _norm(text: str) -> str:
    """去除多余空白，把全角空格折叠。"""
    if not text:
        return ""
    return re.sub(r"\s+", "", text).strip()


def _is_label(text: str) -> bool:
    if not text:
        return False
    norm = _norm(text)
    if not norm:
        return False
    if norm in {_norm(x) for x in _INLINE_LABELS}:
        return True
    if norm in {_norm(x) for x in _SECTION_LABELS}:
        return True
    return False


def _section_of(text: str) -> str:
    if not text:
        return ""
    norm = _norm(text)
    for key, section in _SECTION_LABELS.items():
        if _norm(key) == norm:
            return section
    return ""


def _label_of(text: str) -> str:
    norm = _norm(text)
    for key in _INLINE_LABELS:
        if _norm(key) == norm:
            return key
    return ""


def _dedup_adjacent(cells) -> List[Tuple[object, str]]:
    """按 `_tc` 标识去重（同一合并单元格的多次出现合并为一次）。
    数字逐列存储时各 cell 是独立 `_tc`，不会误吞。"""
    out = []
    prev_tc = None
    for cell in cells:
        if cell._tc is prev_tc:
            continue
        out.append((cell, cell.text.strip()))
        prev_tc = cell._tc
    return out


def _checked_options(cell) -> List[str]:
    """提取 cell 中所有 w:sym=F0FE 处后续的选项文字。"""
    out = []
    for para in cell.paragraphs:
        runs = para.runs
        for i, run in enumerate(runs):
            sym = run._element.find(qn('w:sym'))
            if sym is None:
                continue
            char = sym.get(qn('w:char'))
            if char not in ('F0FE', '00FE'):
                continue
            parts = []
            for j in range(i + 1, min(i + 4, len(runs))):
                if runs[j]._element.find(qn('w:sym')) is not None:
                    break
                txt = runs[j].text.strip()
                if txt:
                    parts.append(txt)
            if parts:
                out.append(''.join(parts))
    return out


def scan_label_pairs(doc) -> Dict[str, str]:
    """扫描全文表格，按 section+label 构造键值对。

    section 仅在同一合并单元格（同 `_tc`）作为行首时持续；新 `_tc` 出现就重置。
    """
    pairs: Dict[str, str] = {}

    for t in doc.tables:
        current_section = ""
        section_tc = None
        for row in t.rows:
            if not row.cells:
                continue
            dedup = _dedup_adjacent(row.cells)
            if not dedup:
                continue

            first_cell, first_text = dedup[0]
            # section 上下文：行首 _tc 仍是 section_tc 则保持；否则按行首判断
            if section_tc is not None and first_cell._tc is section_tc:
                start = 1
            else:
                sec = _section_of(first_text)
                if sec:
                    current_section = sec
                    section_tc = first_cell._tc
                    start = 1
                else:
                    current_section = ""
                    section_tc = None
                    start = 0

            i = start
            while i < len(dedup):
                _, text = dedup[i]
                label = _label_of(text)
                if not label:
                    i += 1
                    continue
                # 收集到下一个 label 之前的文本
                value_parts = []
                j = i + 1
                while j < len(dedup):
                    _, next_text = dedup[j]
                    if _is_label(next_text):
                        break
                    if next_text:
                        value_parts.append(next_text)
                    j += 1
                value = ''.join(value_parts).strip()
                key = f"{current_section}.{label}" if current_section else label
                if value and key not in pairs:
                    pairs[key] = value
                i = j

    return pairs


def scan_checked_in_row(row) -> List[str]:
    """收集行内所有勾选项的文字。"""
    out = []
    for cell in row.cells:
        out.extend(_checked_options(cell))
    return out


def _parse_address(raw: str):
    """从带标签注释的地址串解析省/市/县/详细地址。

    旧版常见形如：
      '山西省省(自治区、直辖市) 运城市地(区、市、州、盟) 河津市县(区、市、旗) 详细地址'
    先把 `(...)` 括号注释、独立的"省/市/县/区"标签词剔除，再按顺序匹配。
    """
    raw = (raw or "").strip()
    if not raw:
        return "", "", "", ""
    # 去除标签括号
    cleaned = re.sub(r"[(（][^()（）]*[)）]", " ", raw)
    cleaned = cleaned.replace("\n", " ").strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    # 拆字段
    province = city = county = ""
    m = re.search(r"([一-龥]+?(?:省|自治区|市))(?=[一-龥]|$|\s)", cleaned)
    if m:
        province = m.group(1)
        cleaned = cleaned[m.end():].strip()
    m = re.search(r"([一-龥]+?(?:市|地区|州|盟))(?=[一-龥]|$|\s)", cleaned)
    if m:
        city = m.group(1)
        cleaned = cleaned[m.end():].strip()
    m = re.search(r"([一-龥]+?(?:县|区|市|旗))(?=[一-龥]|$|\s)", cleaned)
    if m:
        county = m.group(1)
        cleaned = cleaned[m.end():].strip()
    detail = cleaned.strip()
    detail = re.sub(r'^[县区市旗、，,\s]+', '', detail).strip()
    return province, city, county, detail


def _extract_level(text: str) -> str:
    for lv in ("第五级", "第四级", "第三级", "第二级", "第一级"):
        if lv in (text or ""):
            return lv
    return ""


def _leading_code(text: str) -> str:
    """从勾选文字提取前缀编码（'4县（区、市、旗）' → '4'）。"""
    m = re.match(r'^\s*(\d+)', text or "")
    return m.group(1) if m else ""


def read_beian_legacy(doc) -> ProjectData:
    """通用旧版备案表读取（标签扫描）。"""
    data = ProjectData()
    pairs = scan_label_pairs(doc)

    u = data.unit
    u.unit_name = pairs.get("单位名称", "")
    u.credit_code = pairs.get("统一社会信用代码", "") or pairs.get("信用代码", "")
    addr = pairs.get("单位地址", "")
    u.province, u.city, u.county, u.address = _parse_address(addr)
    u.postal_code = pairs.get("邮政编码", "")
    u.admin_code = pairs.get("行政区划代码", "")

    # 负责人
    u.leader = ContactInfo(
        name=pairs.get("leader.姓 名", "") or pairs.get("leader.姓名", ""),
        title=pairs.get("leader.职务/职称", "") or pairs.get("leader.职务", ""),
        office_phone=pairs.get("leader.办公电话", ""),
        mobile=pairs.get("leader.移动电话", ""),
        email=pairs.get("leader.电子邮件", ""),
    )
    u.security_dept = pairs.get("责任部门", "") or pairs.get("网络安全管理机构", "")
    u.security_contact = ContactInfo(
        name=pairs.get("sec_contact.姓 名", "") or pairs.get("sec_contact.姓名", ""),
        title=pairs.get("sec_contact.职务/职称", ""),
        office_phone=pairs.get("sec_contact.办公电话", ""),
        mobile=pairs.get("sec_contact.移动电话", ""),
        email=pairs.get("sec_contact.电子邮件", ""),
    )
    u.data_dept = pairs.get("数据安全管理部门", "")
    u.data_contact = ContactInfo(
        name=pairs.get("data_contact.姓 名", "") or pairs.get("data_contact.姓名", ""),
        title=pairs.get("data_contact.职务/职称", ""),
        office_phone=pairs.get("data_contact.办公电话", ""),
        mobile=pairs.get("data_contact.移动电话", ""),
        email=pairs.get("data_contact.电子邮件", ""),
    )

    # 勾选项：隶属关系/单位类型/行业类别 — 在标签所在 cell 的同行的下一个 cell
    for t in doc.tables:
        for row in t.rows:
            first = row.cells[0].text.strip() if row.cells else ""
            if not first:
                continue
            checked = scan_checked_in_row(row)
            if not checked:
                continue
            code = _leading_code(checked[0])
            if first.startswith("隶属关系") and not u.affiliation:
                u.affiliation = code or checked[0]
            elif first.startswith("单位类型") and not u.unit_type:
                u.unit_type = code or checked[0]
            elif first.startswith("行业类别") and not u.industry:
                u.industry = code or checked[0]

    # 定级对象
    tgt = data.target
    tgt.name = pairs.get("系统名称", "") or pairs.get("定级对象名称", "")
    tgt.code = pairs.get("系统编号", "") or pairs.get("定级对象编号", "")
    tgt.biz_desc = pairs.get("业务描述", "")
    tgt.run_date = pairs.get("运行时间", "") or pairs.get("投入运行时间", "")
    tgt.parent_system = pairs.get("上级系统名称", "")
    tgt.parent_unit = pairs.get("上级系统所属单位", "")

    # 业务类型 / 服务范围 / 等的勾选
    field_label_map = [
        ("业务类型", "biz_type"),
        ("服务范围", "service_scope"),
        ("服务对象", "service_target"),
        ("覆盖范围", "deploy_scope"),
        ("部署范围", "deploy_scope"),
        ("网络性质", "network_type"),
        ("系统互联情况", "interconnect"),
        ("网络互联情况", "interconnect"),
        ("是否分系统", "is_subsystem"),
    ]
    for t in doc.tables:
        for row in t.rows:
            row_text = " ".join(c.text.strip() for c in row.cells)
            for label, attr in field_label_map:
                if label not in row_text:
                    continue
                checked = scan_checked_in_row(row)
                if checked and not getattr(tgt, attr, ""):
                    setattr(tgt, attr, _leading_code(checked[0]) or checked[0])

    # 等级与流程
    g = data.grading
    level_text = pairs.get("信息系统安全保护等级", "") or pairs.get("等级保护对象安全保护等级", "")
    if level_text:
        g.final_level = _extract_level(level_text) or g.final_level
        g.biz_level = g.final_level
        g.service_level = g.final_level
    # 业务/系统服务等级若在表中独立勾选，再校正
    for t in doc.tables:
        for row in t.rows:
            row_text = " ".join(c.text.strip() for c in row.cells)
            checked = scan_checked_in_row(row)
            if "业务信息安全保护等级" in row_text and checked:
                lv = _extract_level(checked[0])
                if lv:
                    g.biz_level = lv
            if "系统服务安全保护等级" in row_text and checked:
                lv = _extract_level(checked[0])
                if lv:
                    g.service_level = lv
            if "信息系统安全保护等级" in row_text and checked:
                lv = _extract_level(checked[0])
                if lv:
                    g.final_level = lv

    g.grading_date = pairs.get("定级时间", "")
    fill_value = pairs.get("填表人", "")
    if fill_value:
        g.filler = fill_value.replace("填表人", "").replace("：", "").strip()
    fill_date = pairs.get("填表日期", "")
    if fill_date:
        g.fill_date = fill_date.replace("填表日期", "").replace("：", "").strip()

    # 专家评审
    for t in doc.tables:
        for row in t.rows:
            row_text = " ".join(c.text.strip() for c in row.cells)
            checked = scan_checked_in_row(row)
            if "专家评审" in row_text or "评审情况" in row_text:
                for c in checked:
                    if "已评审" in c:
                        g.has_review = True
                    elif "未评审" in c:
                        g.has_review = False
            if "系统定级报告" in row_text or "定级报告" in row_text and "附件" in row_text:
                for c in checked:
                    if c.strip() == "有":
                        g.has_report = True
                    elif c.strip() == "无":
                        g.has_report = False
            if "是否有主管部门" in row_text:
                for c in checked:
                    if c.strip() == "有":
                        g.has_supervisor = True
                    elif c.strip() == "无":
                        g.has_supervisor = False
            if "主管部门审批" in row_text or "审核情况" in row_text:
                for c in checked:
                    if "已审" in c:
                        g.supervisor_reviewed = True
                        g.supervisor_review_status = "已审核"
                    elif "未审" in c:
                        g.supervisor_reviewed = False
                        g.supervisor_review_status = "未审核"

    g.supervisor_name = pairs.get("主管部门名称", "")

    # 大数据 / 物联网 / 工控 / 云 — 旧版不强制存在，按标签兜底扫描
    _scan_scenario_legacy(doc, data)

    data.project_name = tgt.name or u.unit_name
    return data


def _scan_scenario_legacy(doc, data):
    """旧版备案表通用应用场景扫描：按行标签匹配是否采用XX技术 / 各系统组成多选。

    不依赖固定行号，遍历所有表格行匹配关键标签；
    旧文档若无该字段，则不写入对应 scenario，保持默认。
    """
    sc = data.scenario
    for t in doc.tables:
        for row in t.rows:
            dedup = _dedup_adjacent(row.cells)
            if not dedup:
                continue
            label_text = ''.join(c.text for c in row.cells if c.text)
            row_label = dedup[0][1] if dedup else ''
            checked = scan_checked_in_row(row)
            if not row_label:
                continue
            # 大数据
            if '是否采用大数据' in row_label and checked:
                if any('是' == c.strip() for c in checked):
                    sc.bigdata.enabled = True
            elif '大数据系统组成' in row_label or '大数据组成' in row_label:
                if checked:
                    sc.bigdata.composition = ','.join(checked)
                    sc.bigdata.enabled = True
            elif '大数据出境' in row_label and checked:
                sc.bigdata.cross_border = checked[0]
            # 云
            elif '是否采用云计算' in row_label and checked:
                if any('是' == c.strip() for c in checked):
                    sc.cloud.enabled = True
            # 移动互联
            elif '是否采用移动互联' in row_label and checked:
                if any('是' == c.strip() for c in checked):
                    sc.mobile.enabled = True
            # 物联网
            elif '是否为物联网' in row_label and checked:
                if any('是' == c.strip() for c in checked):
                    sc.iot.enabled = True
            elif '系统感知层' in row_label and checked:
                sc.iot.perception = ','.join(checked)
                sc.iot.enabled = True
            elif ('系统网络传输层' in row_label or '系统传输层' in row_label) and checked:
                sc.iot.transport = ','.join(checked)
                sc.iot.enabled = True
            # 工控
            elif '是否为工业控制' in row_label and checked:
                if any('是' == c.strip() for c in checked):
                    sc.ics.enabled = True
            elif '系统功能层次' in row_label and checked:
                sc.ics.function_layer = ','.join(checked)
                sc.ics.enabled = True
            elif '工业控制系统组成' in row_label and checked:
                sc.ics.composition = ','.join(checked)
                sc.ics.enabled = True


def read_report_legacy(doc) -> ReportInfo:
    """旧版定级报告通用读取（按段落抓取章节文本）。"""
    report = ReportInfo()
    paragraphs = doc.paragraphs

    for p in paragraphs:
        if p.text.strip():
            report.system_name = p.text.strip()
            break

    section_keys = [
        ("（一）", ("责任主体", "安全责任"), "responsibility"),
        ("（二）", ("基本要素", "网络结构", "定级对象构成"), "composition"),
        ("（三）", ("业务", "承载"), "business_desc"),
        ("（四）", ("安全责任",), "security_resp"),
        ("1、业务信息描述", (), "biz_info_desc"),
        ("2、业务信息受到破坏时所侵害客体", (), "biz_victim"),
        ("3、业务信息受到破坏时", (), "biz_degree"),
        ("1、系统服务描述", (), "svc_desc"),
        ("2、系统服务受到破坏时所侵害客体", (), "svc_victim"),
        ("3、系统服务受到破坏时", (), "svc_degree"),
    ]

    section_buf: Dict[str, List[str]] = {}
    current = ""
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        switched = False
        for prefix, keywords, attr in section_keys:
            if text.startswith(prefix) and (not keywords or any(k in text for k in keywords)):
                current = attr
                switched = True
                break
        if switched:
            continue
        # 顶级章节切换：清空当前段
        if re.match(r"^[一二三四五六七八九十]+、", text):
            current = ""
            continue
        if current:
            section_buf.setdefault(current, []).append(text)

    for attr, lines in section_buf.items():
        setattr(report, attr, "\n".join(lines))

    for p in paragraphs:
        text = p.text.strip()
        if "最终" in text and "等级为" in text:
            lv = _extract_level(text)
            if lv:
                report.final_level = lv
        if "业务信息安全" in text and "等级为" in text:
            lv = _extract_level(text)
            if lv:
                report.biz_level = lv
        if "系统服务安全" in text and "等级为" in text:
            lv = _extract_level(text)
            if lv:
                report.svc_level = lv

    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        header = table.rows[0].cells[0].text.strip()
        if "序号" in header:
            for i in range(1, len(table.rows)):
                row = table.rows[i]
                cells = row.cells
                if len(cells) < 2:
                    continue
                sub = SubSystem(
                    index=cells[0].text.strip(),
                    name=cells[1].text.strip(),
                    description=cells[2].text.strip() if len(cells) > 2 else "",
                )
                if sub.name:
                    report.subsystems.append(sub)
            break

    return report

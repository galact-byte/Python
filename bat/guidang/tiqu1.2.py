import os
import zipfile
import shutil
from pathlib import Path
from docx import Document
import re
import win32com.client
import pythoncom


def extract_project_info_from_docx(docx_path):
    """
    从Word文档的表格中提取项目名称和项目编号
    """
    try:
        # 首先尝试用python-docx读取
        try:
            doc = Document(docx_path)
            result = extract_from_docx_tables(doc)
            print(f"  python-docx提取结果: 项目名称='{result[0]}', 项目编号='{result[1]}'")
            if result[0] and result[1]:
                return result
        except Exception as e:
            print(f"  python-docx读取失败: {e}")

        # 如果python-docx失败，尝试用win32com
        try:
            result = extract_from_doc_win32(docx_path)
            print(f"  win32com提取结果: 项目名称='{result[0]}', 项目编号='{result[1]}'")
            if result[0] and result[1]:
                return result
        except Exception as e:
            print(f"  win32com读取失败: {e}")

        return None, None

    except Exception as e:
        print(f"读取Word文档时出错 {docx_path}: {e}")
        return None, None


def extract_from_docx_tables(doc):
    """
    从docx文档的文本中提取项目信息（非表格格式）
    """
    project_name = None
    project_code = None

    # 获取文档的所有文本内容
    full_text = ""
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + "\n"

    # 也检查表格中的文本（以防信息在表格中）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + "\n"

    print(f"  文档全文前500字符: {repr(full_text[:500])}")

    # 使用正则表达式提取项目名称和项目编号
    # 匹配 "项目名称：" 后面的内容
    name_pattern = r'项目名称[：:]\s*([^\n\r]*?)(?=\s*项目编号|$)'
    name_match = re.search(name_pattern, full_text, re.MULTILINE | re.DOTALL)
    if name_match:
        project_name = name_match.group(1).strip()
        # 移除可能的markdown格式标记
        project_name = re.sub(r'\*\*([^*]*)\*\*', r'\1', project_name)
        print(f"  找到项目名称: {project_name}")
    else:
        print(f"  未找到项目名称，使用的正则: {name_pattern}")

    # 匹配 "项目编号：" 后面的内容
    code_pattern = r'项目编号[：:]\s*([^\n\r]*?)(?=\s*\*\*|$)'
    code_match = re.search(code_pattern, full_text, re.MULTILINE | re.DOTALL)
    if code_match:
        project_code = code_match.group(1).strip()
        # 移除可能的markdown格式标记
        project_code = re.sub(r'\*\*([^*]*)\*\*', r'\1', project_code)
        print(f"  找到项目编号: {project_code}")
    else:
        print(f"  未找到项目编号，使用的正则: {code_pattern}")

    return project_name, project_code


def extract_from_doc_win32(doc_path):
    """
    使用win32com读取.doc文件（文本格式）
    """
    try:
        # 初始化COM
        pythoncom.CoInitialize()

        # 创建Word应用
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False

        # 打开文档
        doc = word_app.Documents.Open(doc_path)

        # 获取文档的全部文本内容
        full_text = doc.Content.Text

        project_name = None
        project_code = None

        # 使用正则表达式提取项目名称和项目编号
        # 匹配 "项目名称：" 后面的内容
        name_pattern = r'项目名称[：:]\s*([^\n\r]*?)(?=\s*项目编号|$)'
        name_match = re.search(name_pattern, full_text)
        if name_match:
            project_name = name_match.group(1).strip()
            # 移除可能的markdown格式标记
            project_name = re.sub(r'\*\*([^*]*)\*\*', r'\1', project_name)

        # 匹配 "项目编号：" 后面的内容
        code_pattern = r'项目编号[：:]\s*([^\n\r]*?)(?=\s*\*\*|$)'
        code_match = re.search(code_pattern, full_text)
        if code_match:
            project_code = code_match.group(1).strip()
            # 移除可能的markdown格式标记
            project_code = re.sub(r'\*\*([^*]*)\*\*', r'\1', project_code)

        # 关闭文档和应用
        doc.Close()
        word_app.Quit()
        pythoncom.CoUninitialize()

        return project_name, project_code

    except Exception as e:
        try:
            if 'doc' in locals():
                doc.Close()
            if 'word_app' in locals():
                word_app.Quit()
            pythoncom.CoUninitialize()
        except:
            pass
        raise e


def clean_filename(filename):
    """
    清理文件名，移除不合法的字符
    """
    # 移除或替换Windows文件名中不允许的字符
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')

    # 移除多余的空格并限制长度
    filename = ' '.join(filename.split())
    if len(filename) > 200:  # 留一些空间给扩展名
        filename = filename[:200]

    return filename


def find_word_file(directory_path, target_subdir="1测评准备", target_filename="07测评工具清单"):
    """
    在指定目录中查找Word文件
    """
    target_dir = os.path.join(directory_path, target_subdir)

    if not os.path.exists(target_dir):
        print(f"目标路径不存在: {target_dir}")
        return None

    # 查找指定名称的Word文件（可能是.doc或.docx格式）
    for file in os.listdir(target_dir):
        filename_without_ext = os.path.splitext(file)[0]
        if filename_without_ext == target_filename and file.lower().endswith(('.docx', '.doc')):
            return os.path.join(target_dir, file)

    return None


def create_zip(source_dir, zip_path):
    """
    创建zip压缩文件
    """
    try:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(source_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    # 计算相对路径
                    arcname = os.path.relpath(file_path, source_dir)
                    zipf.write(file_path, arcname)
        return True
    except Exception as e:
        print(f"创建压缩文件时出错: {e}")
        return False


def batch_process_directories(base_path):
    """
    批量处理目录
    """
    base_path = Path(base_path)

    if not base_path.exists():
        print(f"基础路径不存在: {base_path}")
        return

    processed_count = 0
    failed_count = 0

    # 遍历基础路径下的所有子目录
    for item in base_path.iterdir():
        if item.is_dir():
            print(f"\n正在处理目录: {item.name}")

            # 查找Word文件
            word_file = find_word_file(str(item))

            if not word_file:
                print(f"  未找到Word文件，跳过")
                failed_count += 1
                continue

            print(f"  找到Word文件: {word_file}")

            # 提取项目信息
            project_name, project_code = extract_project_info_from_docx(word_file)

            if not project_name or not project_code:
                print(f"  无法提取项目信息，跳过")
                failed_count += 1
                continue

            print(f"  项目名称: {project_name}")
            print(f"  项目编号: {project_code}")

            # 创建新的目录名
            new_dir_name = f"{project_code}-{project_name}归档资料"
            new_dir_name = clean_filename(new_dir_name)

            # 构建新路径
            new_dir_path = item.parent / new_dir_name

            # 如果新路径已存在，添加序号
            counter = 1
            original_new_dir_path = new_dir_path
            while new_dir_path.exists():
                new_dir_path = original_new_dir_path.parent / f"{original_new_dir_path.name}_{counter}"
                counter += 1

            try:
                # 重命名目录
                item.rename(new_dir_path)
                print(f"  目录重命名为: {new_dir_path.name}")

                # 创建zip文件
                zip_path = new_dir_path.with_suffix('.zip')
                if create_zip(str(new_dir_path), str(zip_path)):
                    print(f"  已创建压缩文件: {zip_path.name}")

                    # 询问是否删除原目录
                    # 为了安全起见，这里不自动删除，用户可以手动删除
                    print(f"  原目录保留在: {new_dir_path}")

                    processed_count += 1
                else:
                    print(f"  压缩失败")
                    failed_count += 1

            except Exception as e:
                print(f"  处理失败: {e}")
                failed_count += 1

    print(f"\n处理完成！")
    print(f"成功处理: {processed_count} 个目录")
    print(f"处理失败: {failed_count} 个目录")


def main():
    """
    主函数
    """
    print("批量目录重命名和压缩工具")
    print("=" * 50)

    # 获取用户输入的基础路径
    base_path = input("请输入要处理的基础目录路径: ").strip()

    if not base_path:
        print("路径不能为空")
        return

    # 移除路径两端的引号（如果有）
    base_path = base_path.strip('"\'')

    # 确认处理
    print(f"\n即将处理路径: {base_path}")
    print("将会:")
    print("1. 在每个子目录中查找 '1测评准备' 文件夹下名为 '07测评工具清单' 的Word文件（.doc或.docx）")
    print("2. 提取项目名称和项目编号")
    print("3. 重命名目录为 '项目编号-项目名称归档资料'")
    print("4. 创建对应的zip压缩文件")

    confirm = input("\n确定要继续吗？(y/N): ").strip().lower()

    if confirm in ['y', 'yes', '是']:
        batch_process_directories(base_path)
    else:
        print("操作已取消")


if __name__ == "__main__":
    main()
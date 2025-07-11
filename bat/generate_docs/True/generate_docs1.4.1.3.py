import os
import re
import shutil
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_user_input():
    """获取用户输入的路径信息"""
    print("请输入路径信息:")
    print("=" * 40)

    # 获取输入路径
    while True:
        print("输入路径选项:")
        print("1. 手动输入路径")
        print("2. 使用当前目录")
        choice = input("请选择 (1/2): ").strip()

        if choice == "1":
            input_path = input("请输入等保完结单文件所在的目录路径: ").strip()
            if not input_path:
                print("路径不能为空，请重新输入")
                continue
            input_path = Path(input_path)
        elif choice == "2":
            input_path = Path.cwd()
            print(f"使用当前目录: {input_path}")
        else:
            print("无效选择，请重新输入")
            continue

        if not input_path.exists():
            print(f"输入路径不存在: {input_path}")
            continue

        if not input_path.is_dir():
            print(f"输入路径不是目录: {input_path}")
            continue

        break

    # 获取输出路径
    while True:
        print("\n输出路径选项:")
        print("1. 手动输入路径")
        print("2. 使用输入路径下的'输出文件'文件夹")
        print("3. 使用当前目录下的'输出文件'文件夹")
        choice = input("请选择 (1/2/3): ").strip()

        if choice == "1":
            output_path = input("请输入输出目录路径: ").strip()
            if not output_path:
                print("路径不能为空，请重新输入")
                continue
            output_path = Path(output_path)
        elif choice == "2":
            output_path = input_path / "输出文件"
            print(f"使用输出路径: {output_path}")
        elif choice == "3":
            output_path = Path.cwd() / "输出文件"
            print(f"使用输出路径: {output_path}")
        else:
            print("无效选择，请重新输入")
            continue

        break

    return input_path, output_path


def check_command_line_args():
    """检查命令行参数"""
    if len(sys.argv) >= 2:
        # 如果有命令行参数，第一个参数作为输入路径
        input_path = Path(sys.argv[1])
        if input_path.exists() and input_path.is_dir():
            print(f"检测到命令行参数，使用输入路径: {input_path}")

            # 第二个参数作为输出路径（如果提供）
            if len(sys.argv) >= 3:
                output_path = Path(sys.argv[2])
            else:
                output_path = input_path / "输出文件"

            print(f"使用输出路径: {output_path}")
            return input_path, output_path

    return None, None


def extract_project_name(filename):
    """
    从文件名中提取项目名称和系统描述，支持中英文括号混排，保留原始括号格式
    """
    # 匹配等保完结单（项目名）（系统描述）或等保完结单(项目名)(系统描述)
    pattern1 = r"等保完结单[（(](.+?)[）)][（(](.+?)[）)]\.docx"
    match = re.match(pattern1, filename)
    if match:
        return match.group(1), match.group(2)

    # 匹配等保完结单（项目名）或等保完结单(项目名)
    pattern2 = r"等保完结单[（(](.+?)[）)]\.docx"
    match = re.match(pattern2, filename)
    if match:
        return match.group(1), None

    return None, None


def extract_project_info_from_docx(docx_path):
    """从等保完结单中提取项目编号和项目名称"""
    try:
        doc = Document(docx_path)

        # 遍历所有表格
        for table in doc.tables:
            # 遍历表格的所有行
            for row in table.rows:
                # 检查每个单元格
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()

                    # 查找包含"项目编号"的单元格
                    if "项目编号" in cell_text:
                        # 项目编号通常在同一行的后面单元格
                        for j in range(i + 1, len(row.cells)):
                            project_code = row.cells[j].text.strip()
                            if project_code and project_code != "项目编号":
                                # 查找项目名称
                                for k in range(len(row.cells)):
                                    if "项目名称" in row.cells[k].text:
                                        for l in range(k + 1, len(row.cells)):
                                            project_name = row.cells[l].text.strip()
                                            if project_name and project_name != "项目名称":
                                                return project_code, project_name
                                break

                # 检查是否有连续的项目编号和项目名称
                if len(row.cells) >= 4:
                    cells_text = [cell.text.strip() for cell in row.cells]

                    # 查找模式：项目编号 | 编号值 | 项目名称 | 名称值
                    for i in range(len(cells_text) - 3):
                        if ("项目编号" in cells_text[i] and
                                "项目名称" in cells_text[i + 2] and
                                cells_text[i + 1] and cells_text[i + 3]):
                            return cells_text[i + 1], cells_text[i + 3]

        # 如果在表格中没找到，尝试在段落中查找
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "项目编号" in text and "项目名称" in text:
                # 使用正则表达式提取
                code_match = re.search(r'项目编号[：:\s]*([^\s]+)', text)
                name_match = re.search(r'项目名称[：:\s]*([^\s]+)', text)
                if code_match and name_match:
                    return code_match.group(1), name_match.group(1)

        return None, None

    except Exception as e:
        print(f"   读取文档 {docx_path} 时出错: {e}")
        return None, None


def get_cell_format(cell):
    """获取单元格的格式信息"""
    format_info = {}

    # 获取第一个段落的格式（如果存在）
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        format_info['alignment'] = paragraph.alignment

        # 获取第一个run的格式（如果存在）
        if paragraph.runs:
            run = paragraph.runs[0]
            format_info['font_name'] = run.font.name
            format_info['font_size'] = run.font.size
            format_info['bold'] = run.font.bold
            format_info['italic'] = run.font.italic
            format_info['underline'] = run.font.underline

    return format_info


def apply_cell_format(cell, format_info, text):
    """应用格式到单元格"""
    # 清空原内容
    cell.text = text

    # 应用格式
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]

        # 应用段落对齐方式
        if 'alignment' in format_info and format_info['alignment']:
            paragraph.alignment = format_info['alignment']

        # 应用字体格式
        if paragraph.runs:
            run = paragraph.runs[0]

            if 'font_name' in format_info and format_info['font_name']:
                run.font.name = format_info['font_name']

            if 'font_size' in format_info and format_info['font_size']:
                run.font.size = format_info['font_size']

            if 'bold' in format_info and format_info['bold'] is not None:
                run.font.bold = format_info['bold']

            if 'italic' in format_info and format_info['italic'] is not None:
                run.font.italic = format_info['italic']

            if 'underline' in format_info and format_info['underline'] is not None:
                run.font.underline = format_info['underline']


def update_template_with_project_info(template_path, output_path, project_code, project_name, source_docx_path):
    """更新模板文件中的项目信息"""
    try:
        # 从源文档获取格式信息
        source_doc = Document(source_docx_path)
        source_format_code = None
        source_format_name = None

        # 从源文档的第一个表格获取格式
        if source_doc.tables:
            first_table = source_doc.tables[0]
            for row in first_table.rows:
                for cell in row.cells:
                    if "项目编号" in cell.text and source_format_code is None:
                        source_format_code = get_cell_format(cell)
                    elif "项目名称" in cell.text and source_format_name is None:
                        source_format_name = get_cell_format(cell)

        # 打开模板文档
        doc = Document(template_path)

        # 处理文档中的段落文本
        for paragraph in doc.paragraphs:
            full_text = paragraph.text

            # 处理项目编号
            if "项目编号：" in full_text:
                # 清除段落现有内容
                paragraph.clear()
                # 重新构建段落
                run = paragraph.add_run("项目编号：")
                run.bold = True
                run.add_text(project_code)
                print(f"   ✅ 已更新段落中的项目编号: {project_code}")

            # 处理项目名称
            elif "项目名称：" in full_text:
                # 清除段落现有内容
                paragraph.clear()
                # 重新构建段落
                run = paragraph.add_run("项目名称：")
                run.bold = True
                run.add_text(project_name)
                print(f"   ✅ 已更新段落中的项目名称: {project_name}")

        # 处理表格中的项目信息
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()

                    # 查找包含"项目编号："的单元格
                    if "项目编号：" in cell_text:
                        # 直接在当前单元格中替换
                        new_text = cell_text.replace("项目编号：", f"项目编号：{project_code}")
                        cell.text = new_text
                        print(f"   ✅ 已更新表格中的项目编号: {project_code}")

                    # 查找包含"项目名称："的单元格
                    elif "项目名称：" in cell_text:
                        # 直接在当前单元格中替换
                        new_text = cell_text.replace("项目名称：", f"项目名称：{project_name}")
                        cell.text = new_text
                        print(f"   ✅ 已更新表格中的项目名称: {project_name}")

                    # 处理分离的标签和值的情况
                    elif "项目编号" in cell_text and "：" not in cell_text:
                        # 项目编号可能在下一个单元格
                        if i + 1 < len(row.cells):
                            next_cell = row.cells[i + 1]
                            if not next_cell.text.strip() or next_cell.text.strip() == "":
                                next_cell.text = project_code
                                print(f"   ✅ 已填入项目编号到下一单元格: {project_code}")

                    elif "项目名称" in cell_text and "：" not in cell_text:
                        # 项目名称可能在下一个单元格
                        if i + 1 < len(row.cells):
                            next_cell = row.cells[i + 1]
                            if not next_cell.text.strip() or next_cell.text.strip() == "":
                                next_cell.text = project_name
                                print(f"   ✅ 已填入项目名称到下一单元格: {project_name}")

        # 保存文档
        doc.save(output_path)
        print(f"   📄 已保存文档: {output_path.name}")
        return True

    except Exception as e:
        print(f"   ❌ 处理模板文件时出错: {e}")
        return False


def get_user_input():
    """获取用户输入的路径信息"""
    print("请输入路径信息:")
    print("=" * 40)

    # 获取输入路径
    while True:
        input_path = input("请输入等保完结单文件所在的目录路径: ").strip()
        if not input_path:
            print("路径不能为空，请重新输入")
            continue

        input_path = Path(input_path)
        if not input_path.exists():
            print(f"输入路径不存在: {input_path}")
            continue

        if not input_path.is_dir():
            print(f"输入路径不是目录: {input_path}")
            continue

        break

    # 获取输出路径
    # while True:
    #     output_path = input("请输入输出目录路径: ").strip()
    #     if not output_path:
    #         print("路径不能为空，请重新输入")
    #         continue
    #
    #     output_path = Path(output_path)
    #     break
    output_path_input = input("请输入输出目录路径:").strip()
    if not output_path_input:
        output_path = input_path / "output"
        print(f"未输入,使用默认输出路径:{output_path}")
    else:
        output_path = Path(output_path_input)

    return input_path, output_path


def ensure_output_directory(output_path):
    """确保输出目录存在，如果不存在则创建"""
    try:
        if not output_path.exists():
            output_path.mkdir(parents=True, exist_ok=True)
            print(f"已创建输出目录: {output_path}")
        elif not output_path.is_dir():
            print(f"错误: 输出路径存在但不是目录: {output_path}")
            return False
        return True
    except Exception as e:
        print(f"创建输出目录时出错: {e}")
        return False


def main():
    # 检查命令行参数
    input_dir, output_dir = check_command_line_args()

    # 如果没有命令行参数，则获取用户输入
    if input_dir is None:
        input_dir, output_dir = get_user_input()

    # 确保输出目录存在
    if not ensure_output_directory(output_dir):
        return

    print(f"\n输入目录: {input_dir}")
    print(f"输出目录: {output_dir}")
    print("=" * 40)

    # 查找模板文件
    template_file = None
    for file in input_dir.glob("*.docx"):
        if "测评过程文档清单" in file.name:
            template_file = file
            break

    if not template_file:
        print("未找到测评过程文档清单")
        return

    print(f"找到模板文件: {template_file.name}")

    # 查找所有等保完结单文件
    completion_files = []
    for file in input_dir.glob("*.docx"):
        # 使用正则同时匹配中文和英文括号，且不修改原始文件名
        if re.match(r"等保完结单[（(].+?[）)]([（(].+?[）)])?\.docx", file.name):
            completion_files.append(file)

    if not completion_files:
        print("未找到等保完结单文件")
        return

    print(f"找到 {len(completion_files)} 个等保完结单文件")

    # 处理每个完结单文件
    success_count = 0
    skip_count = 0
    error_count = 0

    for i, completion_file in enumerate(completion_files, 1):
        print(f"\n[{i}/{len(completion_files)}] 处理文件: {completion_file.name}")

        # 提取项目名称和系统描述
        project_name_from_filename, system_description = extract_project_name(completion_file.name)
        if not project_name_from_filename:
            print(f"❌ 无法从文件名提取项目名称: {completion_file.name}")
            error_count += 1
            continue

        print(f"   项目名称(文件名): {project_name_from_filename}")
        if system_description:
            print(f"   系统描述: {system_description}")

        # 从文档中提取项目信息
        project_code, project_name = extract_project_info_from_docx(completion_file)
        if not project_code or not project_name:
            print(f"❌ 无法从文档中提取项目信息: {completion_file.name}")
            error_count += 1
            continue

        print(f"   项目编号(文档): {project_code}")
        print(f"   项目名称(文档): {project_name}")

        # 生成输出文件名，包含系统描述
        if system_description:
            output_filename = f"测评过程文档清单（{project_name_from_filename}）（{system_description}）.docx"
        else:
            output_filename = f"测评过程文档清单（{project_name_from_filename}）.docx"
        output_path = output_dir / output_filename

        # 检查输出文件是否已存在
        if output_path.exists():
            print(f"⚠️  文件已存在，跳过: {output_filename}")
            skip_count += 1
            continue

        # 复制模板并更新内容
        success = update_template_with_project_info(
            template_file,
            output_path,
            project_code,
            project_name,
            completion_file
        )

        if success:
            print(f"✅ 成功创建: {output_filename}")
            success_count += 1
        else:
            print(f"❌ 创建失败: {output_filename}")
            error_count += 1

    # 显示处理结果摘要
    print(f"\n处理结果摘要:")
    print(f"=" * 40)
    print(f"总文件数: {len(completion_files)}")
    print(f"成功处理: {success_count}")
    print(f"跳过文件: {skip_count}")
    print(f"处理失败: {error_count}")

    if success_count > 0:
        print(f"\n✅ 成功生成的文件保存在: {output_dir}")
    if error_count > 0:
        print(f"⚠️  {error_count} 个文件处理失败，请检查文件格式")


if __name__ == "__main__":
    print("等保完结单处理脚本 v2.0")
    print("=" * 40)
    print("功能: 自动处理等保完结单，生成对应的测评过程文档清单")
    print("支持: 命令行参数、手动输入路径、自动创建输出目录")
    print("=" * 40)

    # 检查是否安装了必要的库
    try:
        import docx
    except ImportError:
        print("错误: 未安装必要的库")
        print("请先安装 python-docx 库:")
        print("pip install python-docx")
        exit(1)

    try:
        main()
    except KeyboardInterrupt:
        print("\n\n用户中断操作")
    except Exception as e:
        print(f"\n程序运行时发生错误: {e}")
        print("请检查文件路径和权限设置")

    print("\n处理完成!")

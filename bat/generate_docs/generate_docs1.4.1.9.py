# 从等保完结单中提取项目名称和项目编号整合到一个过程文档清单中，每个占一页，方便打印
import re
import sys
from pathlib import Path
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.table import Table
from docx.text.paragraph import Paragraph


# 获取用户输入
def get_user_input():
    print("请输入路径信息:")
    print("=" * 40)

    # 获取输入路径
    while True:
        print("输入路径选项:")
        print("1. 手动输入路径")
        print("2. 使用当前目录")
        choice = input("请选择 (1/2): ").strip()

        if choice == "1":
            input_path = input("请输入等保完结单文件所在的目录路径: ").strip()
            if not input_path:
                print("路径不能为空，请重新输入")
                continue
            input_path = Path(input_path)
        elif choice == "2":
            input_path = Path.cwd()
            print(f"使用当前目录: {input_path}")
        else:
            print("无效选择，请重新输入")
            continue

        if not input_path.exists():
            print(f"输入路径不存在: {input_path}")
            continue

        if not input_path.is_dir():
            print(f"输入路径不是目录: {input_path}")
            continue

        break

    # 获取输出路径
    while True:
        print("\n输出路径选项:")
        print("1. 手动输入路径")
        print("2. 使用输入路径下的'output'文件夹")
        print("3. 使用当前目录下的'output'文件夹")
        choice = input("请选择 (1/2/3): ").strip()

        if choice == "1":
            output_path = input("请输入输出目录路径: ").strip()
            if not output_path:
                print("路径不能为空，请重新输入")
                continue
            output_path = Path(output_path)
        elif choice == "2":
            output_path = input_path / "output"
            print(f"使用输出路径: {output_path}")
        elif choice == "3":
            output_path = Path.cwd() / "output"
            print(f"使用输出路径: {output_path}")
        else:
            print("无效选择，请重新输入")
            continue

        break

    return input_path, output_path


# 检查命令行参数
def check_command_line_args():
    if len(sys.argv) >= 2:
        # 如果有命令行参数，第一个参数作为输入路径
        input_path = Path(sys.argv[1])
        if input_path.exists() and input_path.is_dir():
            print(f"检测到命令行参数，使用输入路径: {input_path}")

            # 第二个参数作为输出路径（如果提供）
            if len(sys.argv) >= 3:
                output_path = Path(sys.argv[2])
            else:
                output_path = input_path / "输出文件"

            print(f"使用输出路径: {output_path}")
            return input_path, output_path

    return None, None

# 从文件名中提取项目名称和系统描述，支持中英文括号混排，保留原始括号格式
def extract_project_name(filename):
    # 匹配等保完结单（项目名）（系统描述）或等保完结单(项目名)(系统描述)
    pattern1 = r"等保完结单[（(](.+?)[）)][（(](.+?)[）)]\.docx"
    match = re.match(pattern1, filename)
    if match:
        return match.group(1), match.group(2)

    # 匹配等保完结单（项目名）或等保完结单(项目名)
    pattern2 = r"等保完结单[（(](.+?)[）)]\.docx"
    match = re.match(pattern2, filename)
    if match:
        return match.group(1), None

    return None, None


def extract_project_info_from_docx(docx_path):
    """从等保完结单中提取项目编号和项目名称"""
    try:
        doc = Document(docx_path)

        # 遍历所有表格
        for table in doc.tables:
            # 遍历表格的所有行
            for row in table.rows:
                # 检查每个单元格
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()

                    # 查找包含"项目编号"的单元格
                    if "项目编号" in cell_text:
                        # 项目编号通常在同一行的后面单元格
                        for j in range(i + 1, len(row.cells)):
                            project_code = row.cells[j].text.strip()
                            if project_code and project_code != "项目编号":
                                # 查找项目名称
                                for k in range(len(row.cells)):
                                    if "项目名称" in row.cells[k].text:
                                        for l in range(k + 1, len(row.cells)):
                                            project_name = row.cells[l].text.strip()
                                            if project_name and project_name != "项目名称":
                                                return project_code, project_name
                                break

                # 检查是否有连续的项目编号和项目名称
                if len(row.cells) >= 4:
                    cells_text = [cell.text.strip() for cell in row.cells]

                    # 查找模式：项目编号 | 编号值 | 项目名称 | 名称值
                    for i in range(len(cells_text) - 3):
                        if ("项目编号" in cells_text[i] and
                                "项目名称" in cells_text[i + 2] and
                                cells_text[i + 1] and cells_text[i + 3]):
                            return cells_text[i + 1], cells_text[i + 3]

        # 如果在表格中没找到，尝试在段落中查找
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "项目编号" in text and "项目名称" in text:
                # 使用正则表达式提取
                code_match = re.search(r'项目编号[：:\s]*([^\s]+)', text)
                name_match = re.search(r'项目名称[：:\s]*([^\s]+)', text)
                if code_match and name_match:
                    return code_match.group(1), name_match.group(1)

        return None, None

    except Exception as e:
        print(f"   读取文档 {docx_path} 时出错: {e}")
        return None, None


def set_project_info_style(run, text):
    """根据内容为项目信息字段设置样式（黑体/TNR，小四，加粗）"""
    run.font.bold = True
    run.font.size = Pt(12)  # 小四
    if re.search(r'[\u4e00-\u9fff]', text):  # 包含中文字符
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    else:  # 纯英文/数字
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')


def update_template_with_project_info(template_path, project_code, project_name):
    """更新模板文件中的项目信息，并应用正确的字体样式和调整行间距"""
    try:
        doc = Document(template_path)

        # Define placeholders and their replacement content and line spacing needs
        # We'll use a precise point value for line spacing for better control.
        # Estimate character width for 'project_name' to determine if it wraps.
        # A more robust check might involve measuring text, but this is a reasonable heuristic.
        # Assuming ~0.5cm per character for Chinese/full-width chars at 12pt, 16cm line width (approx)
        # 16cm / 0.5cm/char = 32 chars. Let's use a slightly conservative 28-30.
        # This is an estimate and may need fine-tuning based on your template's specific page margins.
        is_project_name_long = len(project_name) > 28  # Adjusted heuristic

        replacements = {
            "项目编号：": ("项目编号：", project_code, False),
            "项目名称：": ("项目名称：", project_name, is_project_name_long)
        }

        # Update project information in paragraphs
        for p in doc.paragraphs:
            for placeholder, (label, value, adjust_line_spacing) in replacements.items():
                if placeholder in p.text:
                    p.clear()  # Clear paragraph to apply new styles cleanly
                    run_label = p.add_run(label)
                    set_project_info_style(run_label, label)
                    run_value = p.add_run(value)
                    set_project_info_style(run_value, value)

                    p_format = p.paragraph_format
                    p_format.space_before = Pt(0)
                    p_format.space_after = Pt(0)  # Crucial: minimize space after these lines

                    if adjust_line_spacing:
                        # For long project names, use a tighter line spacing to fit two lines compactly
                        p_format.line_spacing = Pt(14)  # Adjusted to be tighter, e.g., 14pt or even 13.5pt if needed
                    else:
                        p_format.line_spacing = 1.0  # Standard single line spacing

                    # Add a page break after the project information block if it's not the last project
                    # This ensures each project starts on a new conceptual page block within the merged document.
                    # This is handled in create_merged_document's copy_document_content.
                    break

        # Update project information in tables
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()

                    # Case 1: Label and value in the same cell
                    for placeholder, (label, value, adjust_line_spacing) in replacements.items():
                        if placeholder in cell_text:
                            p = cell.paragraphs[0]
                            p.clear()
                            run_label = p.add_run(label)
                            set_project_info_style(run_label, label)
                            run_value = p.add_run(value)
                            set_project_info_style(run_value, value)

                            p_format = p.paragraph_format
                            p_format.space_before = Pt(0)
                            p_format.space_after = Pt(0)

                            if adjust_line_spacing:
                                p_format.line_spacing = Pt(14)  # Consistent with paragraph updates
                            else:
                                p_format.line_spacing = 1.0  # Or Pt(12) for exact match to font size

                            # Ensure vertical alignment in cell
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            break
                    else:
                        continue  # Continue to next replacement if this one wasn't found
                    break  # Break from inner loop (replacements) to next cell

                # This loop handles Case 2 (label and value in adjacent cells)
                # It needs to be separate or carefully integrated to avoid double processing
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if "项目编号" in cell_text and "：" not in cell_text:
                        p_label = cell.paragraphs[0]
                        p_label.clear()
                        set_project_info_style(p_label.add_run("项目编号"), "项目编号")
                        p_format_label = p_label.paragraph_format
                        p_format_label.space_before = Pt(0)
                        p_format_label.space_after = Pt(0)
                        p_format_label.line_spacing = 1.0
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        if i + 1 < len(row.cells):
                            next_cell = row.cells[i + 1]
                            p_value = next_cell.paragraphs[0]
                            p_value.clear()
                            set_project_info_style(p_value.add_run(project_code), project_code)
                            p_format_value = p_value.paragraph_format
                            p_format_value.space_before = Pt(0)
                            p_format_value.space_after = Pt(0)
                            p_format_value.line_spacing = 1.0
                            next_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    elif "项目名称" in cell_text and "：" not in cell_text:
                        p_label = cell.paragraphs[0]
                        p_label.clear()
                        set_project_info_style(p_label.add_run("项目名称"), "项目名称")
                        p_format_label = p_label.paragraph_format
                        p_format_label.space_before = Pt(0)
                        p_format_label.space_after = Pt(0)
                        p_format_label.line_spacing = 1.0
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                        if i + 1 < len(row.cells):
                            next_cell = row.cells[i + 1]
                            p_value = next_cell.paragraphs[0]
                            p_value.clear()
                            set_project_info_style(p_value.add_run(project_name), project_name)

                            p_format_value = p_value.paragraph_format
                            p_format_value.space_before = Pt(0)
                            p_format_value.space_after = Pt(0)
                            if is_project_name_long:  # Use the same heuristic for cells
                                p_format_value.line_spacing = Pt(14)  # Consistent
                            else:
                                p_format_value.line_spacing = 1.0
                            next_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        return doc

    except Exception as e:
        print(f"   ❌ 处理模板文件时出错: {e}")
        return None


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 细线
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)

    tblPr.append(borders)


def set_font_style(run, text, is_bold=False, is_header=False, is_table=False):
    # 判断中文或英文
    if re.search(r'[\u4e00-\u9fff]', text):  # 中文
        if is_header:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)
            run.font.bold = True
        elif is_table:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(9)  # 小五 (9pt)
        else:  # 正文 (project info)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(12)  # 小四 (12pt)
            run.font.bold = is_bold
    else:  # 英文或数字
        run.font.name = 'Times New Roman' if not is_table else 'Calibri'
        run._element.rPr.rFonts.set(qn('w:ascii'), run.font.name)
        run.font.size = Pt(12 if not is_table else 9)  # 小四 (12pt) for non-table, small_five (9pt) for table
        run.font.bold = is_bold if not is_table else False


def set_title_style(paragraph):
    """设置标题样式：宋体，三号，加粗，居中"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        run = paragraph.runs[0]
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(16)  # 三号
        run.font.bold = True


def copy_document_content(source_doc, target_doc):
    """将源文档的内容（段落和表格）按顺序复制到目标文档，并应用紧凑的布局样式。"""

    # Track if we've added the main title for the current project block
    added_project_info = False

    for element in source_doc.element.body:
        if element.tag.endswith('p'):
            para = Paragraph(element, source_doc)

            # Skip adding extra paragraphs if they are just empty or carry over
            # from hidden elements, unless it's a critical element.
            if not para.text.strip() and not para.runs:
                continue

            # Project Code and Project Name paragraphs
            if "项目编号：" in para.text or "项目名称：" in para.text:
                new_paragraph = target_doc.add_paragraph()
                p_format = new_paragraph.paragraph_format
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)  # Ensure no extra space after project info
                # Line spacing is handled in update_template_with_project_info for these specific paragraphs
                if para.alignment:
                    new_paragraph.alignment = para.alignment

                for run in para.runs:
                    new_run = new_paragraph.add_run(run.text)
                    set_project_info_style(new_run, run.text)  # Use project info style
                added_project_info = True

            # "测评过程文档清单" title
            elif para.text.strip() == "测评过程文档清单":
                new_paragraph = target_doc.add_paragraph("测评过程文档清单")
                p_format = new_paragraph.paragraph_format
                p_format.space_before = Pt(
                    0)  # Crucial: Remove space before the title if it follows project info directly
                p_format.space_after = Pt(18)  # Keep space after title
                p_format.line_spacing = 1.0
                set_title_style(new_paragraph)
            else:
                # Other regular paragraphs, ensure compact spacing
                new_paragraph = target_doc.add_paragraph()
                p_format = new_paragraph.paragraph_format
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(3)  # Small space after regular paragraphs
                p_format.line_spacing = 1.2  # Slightly more relaxed for general text

                if para.alignment:
                    new_paragraph.alignment = para.alignment

                for run in para.runs:
                    new_run = new_paragraph.add_run(run.text)
                    set_font_style(new_run, run.text, is_bold=run.bold)

        elif element.tag.endswith('tbl'):
            table = Table(element, source_doc)

            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0

            if rows > 0 and cols > 0:
                new_table = target_doc.add_table(rows=rows, cols=cols)
                if table.style:
                    new_table.style = table.style
                set_table_border(new_table)

                col_widths = [Cm(2.27), Cm(6.42), Cm(6.32)]
                if len(col_widths) == cols:
                    for i, row in enumerate(new_table.rows):
                        for j, cell in enumerate(row.cells):
                            cell.width = col_widths[j]
                else:
                    print(f"警告: 预设列宽数量 ({len(col_widths)}) 与表格列数 ({cols}) 不匹配，跳过列宽设置。")

                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                            new_cell = new_table.rows[i].cells[j]
                            new_cell.text = ""  # Clear existing text to ensure clean copy

                            # Copy paragraphs from source cell to new cell
                            for k, paragraph in enumerate(cell.paragraphs):
                                # If the source cell has multiple paragraphs, add new ones; otherwise, use the first.
                                new_para = new_cell.paragraphs[0] if k == 0 else new_cell.add_paragraph()
                                p_format = new_para.paragraph_format

                                # Set consistent tight spacing for all table cell paragraphs
                                p_format.space_before = Pt(0)
                                p_format.space_after = Pt(0)
                                p_format.line_spacing = 1.0  # Use single line spacing for better control (or Pt(9) if more precise)

                                # Apply horizontal alignment based on column and row type
                                if i == 0:  # Header row
                                    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    # Header cells should also be vertically centered
                                    new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                else:  # Data rows
                                    if j in [0, 2]:  # '序号' (0th column) and '文档编号' (2nd column)
                                        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    elif j == 1:  # '文档名称' (1st column)
                                        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    # Ensure all data cells are vertically centered
                                    new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                                for run in paragraph.runs:
                                    new_run = new_para.add_run(run.text)
                                    is_header = (i == 0)
                                    set_font_style(new_run, run.text, is_header=is_header, is_table=True)

                for row in new_table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), '400')
                    trHeight.set(qn('w:hRule'), 'atLeast')
                    trPr.append(trHeight)


def create_merged_document(template_path, project_info_list, output_path):
    """创建合并的文档，包含所有项目的完整测评过程文档清单"""
    try:
        merged_doc = Document()

        # For the first project, no page break is needed at the very beginning
        # We will add a page break *before* subsequent projects.
        for i, project_info in enumerate(project_info_list):
            project_code, project_name, _, _ = project_info
            print(f"   正在添加项目 {i + 1}: {project_name}")

            # Add a page break before each project section *except* the very first one
            if i > 0:
                merged_doc.add_page_break()

            # Create an updated document object for the current project
            project_doc = update_template_with_project_info(template_path, project_code, project_name)

            if project_doc:
                # Copy the content of the project document to the merged document
                copy_document_content(project_doc, merged_doc)
                print(f"   ✅ 已添加项目 {i + 1}: {project_name}")
            else:
                print(f"   ❌ 处理项目 {i + 1} 失败: {project_name}")

        # Save the merged document
        merged_doc.save(output_path)
        print(f"\n   📄 已保存合并文档: {output_path.name}")
        return True

    except Exception as e:
        print(f"   ❌ 创建合并文档时出错: {e}")
        return False


def ensure_output_directory(output_path):
    """确保输出目录存在，如果不存在则创建"""
    try:
        if not output_path.exists():
            output_path.mkdir(parents=True, exist_ok=True)
            print(f"已创建输出目录: {output_path}")
        elif not output_path.is_dir():
            print(f"错误: 输出路径存在但不是目录: {output_path}")
            return False
        return True
    except Exception as e:
        print(f"创建输出目录时出错: {e}")
        return False


def main():
    # 检查命令行参数
    input_dir, output_dir = check_command_line_args()

    # 如果没有命令行参数，则获取用户输入
    if input_dir is None:
        input_dir, output_dir = get_user_input()

    # 确保输出目录存在
    if not ensure_output_directory(output_dir):
        return

    print(f"\n输入目录: {input_dir}")
    print(f"输出目录: {output_dir}")
    print("=" * 40)

    # 查找模板文件
    template_file = None
    for file in input_dir.glob("*.docx"):
        if "测评过程文档清单" in file.name:
            template_file = file
            break

    if not template_file:
        print("未找到测评过程文档清单模板文件")
        return

    print(f"找到模板文件: {template_file.name}")

    # 查找所有等保完结单文件
    completion_files = []
    for file in input_dir.glob("*.docx"):
        if re.match(r"等保完结单[（(].+?[）)]([（(].+?[）)])?\.docx", file.name):
            completion_files.append(file)

    if not completion_files:
        print("未找到等保完结单文件")
        return

    print(f"找到 {len(completion_files)} 个等保完结单文件")

    # 收集所有项目信息
    project_info_list = []
    error_count = 0

    for i, completion_file in enumerate(completion_files, 1):
        print(f"\n[{i}/{len(completion_files)}] 处理文件: {completion_file.name}")

        project_name_from_filename, system_description = extract_project_name(completion_file.name)
        if not project_name_from_filename:
            print(f"❌ 无法从文件名提取项目名称: {completion_file.name}")
            error_count += 1
            continue

        print(f"   项目名称(文件名): {project_name_from_filename}")
        if system_description:
            print(f"   系统描述: {system_description}")

        project_code, project_name = extract_project_info_from_docx(completion_file)
        if not project_code or not project_name:
            print(f"❌ 无法从文档中提取项目信息: {completion_file.name}")
            error_count += 1
            continue

        print(f"   项目编号(文档): {project_code}")
        print(f"   项目名称(文档): {project_name}")

        project_info_list.append((project_code, project_name, project_name_from_filename, system_description))
        print(f"✅ 已收集项目信息")

    if not project_info_list:
        print("❌ 没有成功提取到任何项目信息")
        return

    # Generate unified output filename
    output_filename = f"测评过程文档清单({len(project_info_list)}个项目).docx"
    output_path = output_dir / output_filename

    if output_path.exists():
        overwrite = input(f"文件已存在: {output_filename}\n是否覆盖? (y/n): ").strip().lower()
        if overwrite != 'y':
            print("操作取消")
            return

    # Create the merged document
    print(f"\n正在创建合并文档，包含 {len(project_info_list)} 个项目的完整测评过程文档清单...")
    success = create_merged_document(template_file, project_info_list, output_path)

    print(f"\n处理结果摘要:")
    print(f"=" * 40)
    print(f"总文件数: {len(completion_files)}")
    print(f"成功处理: {len(project_info_list)}")
    print(f"处理失败: {error_count}")

    if success:
        print(f"\n✅ 成功创建合并文档: {output_filename}")
        print(f"   文件保存在: {output_dir}")
    else:
        print(f"❌ 创建合并文档失败")

    if error_count > 0:
        print(f"⚠️  {error_count} 个文件处理失败，请检查文件格式")


if __name__ == "__main__":
    print("等保完结单处理脚本 v3.1 - 格式修正版")
    print("=" * 50)
    print("功能: 自动处理等保完结单，生成包含所有项目的测评过程文档清单合集")
    print("特点: 修正标题、项目信息和表格的字体格式")
    print("=" * 50)

    try:
        import docx
    except ImportError:
        print("错误: 未安装必要的库")
        print("请先安装 python-docx 库:")
        print("pip install python-docx")
        sys.exit(1)

    try:
        main()
    except KeyboardInterrupt:
        print("\n\n用户中断操作")
    except Exception as e:
        print(f"\n程序运行时发生错误: {e}")
        print("请检查文件路径和权限设置")

    print("\n处理完成!")
